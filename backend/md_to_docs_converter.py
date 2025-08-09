import os
import re
from datetime import datetime
from pathlib import Path
import pypandoc
from docx import Document as DocxDocument
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docxtpl import DocxTemplate
import logging

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class MarkdownTitleExtractor:
    """Extract title page information from markdown"""
    
    def __init__(self, markdown_content: str):
        self.content = markdown_content
        
    def extract_title_info(self) -> dict:
        """Extract information needed for title page and document information"""
        lines = self.content.split('\n')
        
        title_info = {
            'program_name': '',
            'description': '',
            'ricefw_id': '',
            'module_name': 'Human Resource',
            'document_type': 'Functional Specification Design (FSD)',
            'current_date': datetime.now().strftime('%Y'),
            'file_name': '',
            # New fields for second page
            'project_name': 'System Integrator for Management Information System Towards Single Source of Truth Implementation Program',
            'document_location': '',
            'related_documents': [],
            'reviewers': [],
            'version_history': [],
            'generated_date': datetime.now().strftime('%Y-%m-%d'),
            'generated_time': datetime.now().strftime('%H:%M:%S'),
            # New fields for table of contents
            'table_of_contents': [],
            # New fields for fourth page
            'user_requirements': '',
            'assumptions': [],
            'sap_program_name': '',
            'transaction_code': '',
            'menu_path': '',
            # New fields for DESAIN section
            'selection_screen_table': [],
            'detail_processing_table': [],
            'valid_datasets_table': [],
            'country_info_table': [],
            'report_description': '',
            'authorization_info': '',
            'design_constraints': '',
            # New fields for ERROR HANDLING section
            'error_handling_table': [],
            # New fields for TESTING REQUIREMENTS section
            'testing_requirements_table': [],
            'test_data_location': '',
            'test_transaction': '',
            'test_menu_path': '',
            # ADD THIS NEW FIELD:
            'design_changes': '',  # Add this line
            # Additional design tables
            'valid_datasets_table': [],
            'country_info_table': [],
            'currency_t500c_table': [],
            'currency_t001_table': []
        }
        
        current_section = None
        in_general_requirements = False
        in_existing_sap_objects = False
        in_design_section = False
        in_error_handling = False
        in_testing_requirements = False
        in_design_changes = False
        current_design_subsection = None
        
        for i, line in enumerate(lines):
            line = line.strip()
            
            # Track which major section we're in
            if line.startswith('## ') and ('GENERAL REQUIREMENTS' in line.upper() or 'PERSYARATAN UMUM'    in line.upper() or 'INFORMASI DOKUMEN'   in line.upper()):
                in_general_requirements = True
                in_existing_sap_objects = False
                in_design_section = False
                in_error_handling = False
                in_testing_requirements = False
                current_section = None
                continue
            elif line.startswith('## ') and ('EXISTING SAP OBJECTS' in line.upper() or 'OBJEK SAP EXISTING' in line.upper()):
                in_existing_sap_objects = True
                in_general_requirements = False
                in_design_section = False
                in_error_handling = False
                in_testing_requirements = False
                current_section = None
                continue
            elif line.startswith('## ') and ('DESIGN' in line.upper() or 'DESAIN' in line.upper()):
                in_design_section = True
                in_general_requirements = False
                in_existing_sap_objects = False
                in_error_handling = False
                in_testing_requirements = False
                current_section = None
                continue
            elif line.startswith('## ') and ('ERROR HANDLING' in line.upper() or 'PENANGANAN ERROR' in line.upper()):
                in_error_handling = True
                in_general_requirements = False
                in_existing_sap_objects = False
                in_design_section = False
                in_testing_requirements = False
                current_section = None
                continue
            elif line.startswith('## ') and ('TESTING REQUIREMENTS' in line.upper() or 'PERSYARATAN PENGUJIAN' in line.upper()):
                in_testing_requirements = True
                in_general_requirements = False
                in_existing_sap_objects = False
                in_design_section = False
                in_error_handling = False
                current_section = None
                continue
            elif line.startswith('## ') and ('DESIGN CHANGE' in line.upper() or 'PERUBAHAN DESAIN' in line.upper()):
                in_design_changes = True
                in_testing_requirements = False
                in_general_requirements = False
                in_existing_sap_objects = False
                in_design_section = False
                in_error_handling = False
                current_section = None
                continue
            elif line.startswith('## ') and not line.startswith('### '):
                in_general_requirements = False
                in_existing_sap_objects = False
                in_design_section = False
                in_error_handling = False
                in_testing_requirements = False
                current_section = None
            
            # Track design subsections
            if in_design_section and line.startswith('### '):
                header = line.replace('###','').strip().lower()
                if 'selection screen' in header:
                    current_design_subsection = 'selection_screen'
                elif 'detail processing' in header:
                    current_design_subsection = 'detail_processing'
                elif 'Detail Process Only valid datasets' in line:
                    current_design_subsection = 'valid_datasets'
                elif 'Form Get_Country_Info' in line:
                    current_design_subsection = 'country_info'
                elif 'Form Get_Currency_T500C' in line:
                    current_design_subsection = 'currency_t500c'
                elif 'Form Get_Currency_T001' in line:
                    current_design_subsection = 'currency_t001'
                else:
                    current_design_subsection = None
                continue

            
            # Extract program name from ## heading
            if line.startswith('## ') and not title_info['program_name']:
                title_info['program_name'] = line.replace('## ', '').strip()
            
            # Extract description from ### heading  
            if line.startswith('### ') and not title_info['description']:
                title_info['description'] = line.replace('### ', '').strip()
            
            # Extract RICEFW ID from document or comments
            if 'RICEFW ID' in line or re.search(r'RHR\d+', line):
                match = re.search(r'RHR\d+', line)
                if match:
                    title_info['ricefw_id'] = match.group()
            
            # Extract generated date
            if '**Generated on:**' in line:
                date_match = re.search(r'(\d{4}-\d{2}-\d{2})\s+(\d{2}:\d{2}:\d{2})', line)
                if date_match:
                    title_info['generated_date'] = date_match.group(1)
                    title_info['generated_time'] = date_match.group(2)
                    title_info['generated_datetime'] = f"{date_match.group(1)} {date_match.group(2)}"
                else:
                    date_only_match = re.search(r'(\d{4}-\d{2}-\d{2})', line)
                    if date_only_match:
                        title_info['generated_date'] = date_only_match.group(1)
            
            # Extract project name
            if '**Project**:' in line:
                project_match = re.search(r'\*\*Project\*\*:\s*(.+)', line)
                if project_match:
                    extracted_project = project_match.group(1).strip()
                    if extracted_project and extracted_project != 'System Integrator for Management Information System Towards Single Source of Truth Implementation Program':
                        title_info['project_name'] = extracted_project
            
            # Extract document location
            if '**Document Location**:' in line:
                doc_location_match = re.search(r'\*\*Document Location\*\*:\s*(.+)', line)
                if doc_location_match:
                    title_info['document_location'] = doc_location_match.group(1).strip()
            
            # Extract User Requirements from GENERAL REQUIREMENTS section
            if in_general_requirements and '**User Requirements**:' in line:
                # grab the rest of this line
                user_req = re.sub(r'\*\*User Requirements\*\*:\s*', '', line).strip()
                # now consume any following lines until a blank or next section
                j = i + 1
                while j < len(lines):
                    nxt = lines[j].rstrip()
                    # stop if blank, or we hit Assumptions, or a new "##" section
                    if not nxt or nxt.startswith('**Assumptions**') or nxt.startswith('## '):
                        break
                    user_req += ' ' + nxt.strip()
                    j += 1
                title_info['user_requirements'] = user_req

            
            # Extract Assumptions from GENERAL REQUIREMENTS section
            if in_general_requirements and '**Assumptions**:' in line:
                current_section = 'assumptions'
                continue
            elif in_general_requirements and current_section == 'assumptions' and line.startswith('- '):
                assumption = line.replace('- ', '').strip()
                title_info['assumptions'].append(assumption)
            
            # Extract SAP Objects information from EXISTING SAP OBJECTS section
            if in_existing_sap_objects and '**Program Name**:' in line:
                prog_match = re.search(r'\*\*Program Name\*\*:\s*(.+)', line)
                if prog_match:
                    title_info['sap_program_name'] = prog_match.group(1).strip()
            
            if in_existing_sap_objects and '**Transaction Code**:' in line:
                trans_match = re.search(r'\*\*Transaction Code\*\*:\s*(.+)', line)
                if trans_match:
                    title_info['transaction_code'] = trans_match.group(1).strip()
            
            if in_existing_sap_objects and '**Menu Path**:' in line:
                menu_match = re.search(r'\*\*Menu Path\*\*:\s*(.+)', line)
                if menu_match:
                    title_info['menu_path'] = menu_match.group(1).strip()
            
            # Extract table data from DESIGN section
            if in_design_section and current_design_subsection and line.startswith('|') and '---' not in line:
                # Parse table row
                if current_design_subsection == 'selection_screen':
                    if not any('Parameter' in cell for cell in line.split('|')):  # Skip header row
                        cells = [cell.strip() for cell in line.split('|')[1:-1]]  # Remove empty first/last elements
                        if len(cells) >= 6:
                            title_info['selection_screen_table'].append({
                                'parameter': cells[0],
                                'type': cells[1],
                                'description': cells[2],
                                'mandatory': cells[3],
                                'select_option': cells[4],
                                'no_intervals': cells[5]
                            })
                
                elif current_design_subsection == 'detail_processing':
                    if not any('Field Name' in cell for cell in line.split('|')):  # Skip header row
                        cells = [cell.strip() for cell in line.split('|')[1:-1]]  # Remove empty first/last elements
                        if len(cells) >= 5:
                            title_info['detail_processing_table'].append({
                                'field_name': cells[0],
                                'technical_field': cells[1],
                                'source_table': cells[2],
                                'processing_logic': cells[3],
                                'processing_type': cells[4]
                            })
                elif current_design_subsection in ['valid_datasets', 'country_info', 'currency_t500c', 'currency_t001']:
                    if not any('Data' in cell for cell in line.split('|')):  # Skip header row
                        cells = [cell.strip() for cell in line.split('|')[1:-1]]
                        if len(cells) >= 2:
                            entry = {
                                'data': cells[0],
                                'condition': cells[1]
                            }
                            if current_design_subsection == 'valid_datasets':
                                title_info['valid_datasets_table'].append(entry)
                            elif current_design_subsection == 'country_info':
                                title_info['country_info_table'].append(entry)
                            elif current_design_subsection == 'currency_t500c':
                                title_info['currency_t500c_table'].append(entry)
                            elif current_design_subsection == 'currency_t001':
                                title_info['currency_t001_table'].append(entry)

                # elif current_design_subsection == 'valid_datasets':
                #     # Valid datasets table (2 cols: Data | Kondisi)
                #     cells = [c.strip() for c in line.split('|')[1:-1]]
                #     if len(cells) == 2:
                #         title_info['valid_datasets_table'].append({
                #             'data':      cells[0],
                #             'condition': cells[1],
                #         })

                # elif current_design_subsection == 'country_info':
                #     # Country info table (2 cols: Data | Kondisi)
                #     cells = [c.strip() for c in line.split('|')[1:-1]]
                #     if len(cells) == 2:
                #         title_info['country_info_table'].append({
                #             'data':      cells[0],
                #             'condition': cells[1],
                #         })

            
            # Extract ERROR HANDLING table data
            if in_error_handling and line.startswith('|') and '---' not in line:
                if not any('Error Description' in cell for cell in line.split('|')):  # Skip header row
                    cells = [cell.strip() for cell in line.split('|')[1:-1]]  # Remove empty first/last elements
                    if len(cells) >= 5:
                        title_info['error_handling_table'].append({
                            'no': cells[0],
                            'error_description': cells[1],
                            'resolution': cells[2],
                            'error_code': cells[3],
                            'severity': cells[4]
                        })
            
            # Extract TESTING REQUIREMENTS table data
            if in_testing_requirements and line.startswith('|') and '---' not in line:
                if not any('Test Condition' in cell for cell in line.split('|')):  # Skip header row
                    cells = [cell.strip() for cell in line.split('|')[1:-1]]  # Remove empty first/last elements
                    if len(cells) >= 5:
                        title_info['testing_requirements_table'].append({
                            'no': cells[0],
                            'test_condition': cells[1],
                            'expected_result': cells[2],
                            'test_data': cells[3],
                            'priority': cells[4]
                        })
            
            # Track sections for multi-line parsing (existing code)
            if '**Related Documents**:' in line:
                current_section = 'related_documents'
                continue
            elif '**Reviewers**:' in line:
                current_section = 'reviewers'
                continue
            elif '**Version History**:' in line:
                current_section = 'version_history'
                continue
            elif line.startswith('##') and not line.startswith('### '):
                if current_section not in ['assumptions']:  # Don't reset assumptions in ## sections
                    current_section = None
            
            # Parse multi-line sections (existing code)
            if current_section == 'related_documents' and line.startswith('  - '):
                title_info['related_documents'].append(line.replace('  - ', '').strip())
            elif current_section == 'reviewers' and line.startswith('  - '):
                reviewer_info = line.replace('  - ', '').strip()
                title_info['reviewers'].append(reviewer_info)
            elif current_section == 'version_history' and line.startswith('  - '):
                version_info = line.replace('  - ', '').strip()
                title_info['version_history'].append(version_info)
        
        # If no RICEFW ID found, try to infer from program name pattern
        if not title_info['ricefw_id'] and title_info['program_name']:
            # For ZHR_R_IT0015, try to map to RHR pattern
            if 'IT0015' in title_info['program_name']:
                title_info['ricefw_id'] = 'RHR006'  # Based on your documents, IT0015 maps to RHR006
            elif 'IT0267' in title_info['program_name']:
                title_info['ricefw_id'] = 'RHR018'  # Based on your documents, IT0267 maps to RHR018
            else:
                # Default pattern
                title_info['ricefw_id'] = 'RHR041'
        
        # Create file name from program name and description
        # if title_info['program_name'] and title_info['description']:
        #     title_info['file_name'] = f"Report of {title_info['description']}"
        # elif title_info['description']:
        #     title_info['file_name'] = f"Report of {title_info['description']}"
        if title_info['description']:
            # Use description directly for file name to match document heading
            title_info['file_name'] = title_info['description']
        else:
            title_info['file_name'] = 'Document'
        
        # If document location not found, create it from program name
        if not title_info['document_location'] and title_info['program_name']:
            # Convert ZHR_R_IT0015 to zhr_r_it0015.html
            doc_name = title_info['program_name'].lower().replace('_', '_')
            title_info['document_location'] = f"{doc_name}.html"
        
        # Set default test values based on extracted information
        title_info['test_data_location'] = 'Data uji tersedia di environment development SAP'
        title_info['test_transaction'] = title_info.get('transaction_code', 'N/A')
        title_info['test_menu_path'] = title_info.get('menu_path', 'N/A')
        
        # Extract table of contents from markdown structure
        title_info['table_of_contents'] = self._extract_table_of_contents(lines)
        
        return title_info
    
    def _extract_table_of_contents(self, lines: list) -> list:
        """Extract table of contents from markdown headings"""
        toc_items = []
        page_number = 4  # Start from page 4 (after title, doc info, and TOC pages)
        
        # Define section mappings (English to Indonesian)
        # section_mappings = {
        #     'DOCUMENT INFORMATION': 'INFORMASI DOKUMEN',
        #     'GENERAL REQUIREMENTS': 'PERSYARATAN UMUM',
        #     'EXISTING SAP OBJECTS': 'OBJEK SAP EXISTING YANG TERKAIT DENGAN REPORT (Existing SAP Object Related to the Reports)',
        #     'DESIGN': 'DESAIN',
        #     'Selection Screen': 'Selection Screen',
        #     'Detail Processing': 'Detail Processing',
        #     'ERROR HANDLING': 'PENANGANAN ERROR',
        #     'DESIGN ALTERNATIVES': 'DESAIN ALTERNATIF',
        #     'TESTING REQUIREMENTS': 'PERSYARATAN PENGUJIAN',
        #     'DESIGN CHANGE': 'PERUBAHAN DESAIN'
        # }
        
        section_mappings = {
            'INFORMASI DOKUMEN': 'INFORMASI DOKUMEN',
            'PERSYARATAN UMUM': 'PERSYARATAN UMUM',
            'OBJEK SAP EXISTING YANG TERKAIT DENGAN REPORT (Existing SAP Object Related to the Reports)': 'OBJEK SAP EXISTING YANG TERKAIT DENGAN REPORT (Existing SAP Object Related to the Reports)',
            'DESAIN': 'DESAIN',
            'Selection Screen': 'Selection Screen',
            'Detail Processing': 'Detail Processing',
            'Detail Process Only valid datasets': 'Detail Process Only valid datasets',
            'Form Get_Country_Info': 'Form Get_Country_Info',
            'Form Get_Currency_T500C': 'Form Get_Currency_T500C',
            'Form Get_Currency_T001': 'Form Get_Currency_T001',
            'PENANGANAN ERROR': 'PENANGANAN ERROR',
            'DESAIN ALTERNATIF': 'DESAIN ALTERNATIF',
            'PERSYARATAN PENGUJIAN': 'PERSYARATAN PENGUJIAN',
            'PERUBAHAN DESAIN': 'PERUBAHAN DESAIN'
        }
        
        # Add fixed TOC items first
        toc_items.append({
            'number': '1',
            'title': 'INFORMASI DOKUMEN',
            'page': 2
        })
        toc_items.append({
            'number': '2',
            'title': 'DAFTAR ISI',
            'page': 3
        })
        
        section_counter = 3
        current_main_section = 0
        
        # First pass: collect all sections and subsections
        sections_found = []
        
        for line in lines:
            line = line.strip()
            
            # Process main sections (## headings with numbers like "## 1. DOCUMENT INFORMATION")
            main_section_match = re.match(r'## (\d+)\.\s*(.+)', line)
            if main_section_match:
                section_num = int(main_section_match.group(1))
                original_title = main_section_match.group(2).strip()
                
                # Skip title sections (## ZHR_R_IT0015)
                if section_num >= 1:
                    sections_found.append({
                        'type': 'main',
                        'original_number': section_num,
                        'title': original_title,
                        'subsections': []
                    })
            
            # Process subsections (### headings like "### 4.1 Selection Screen")
            subsection_match = re.match(r'### (\d+)\.(\d+)\s*(.+)', line)
            if subsection_match:
                main_num = int(subsection_match.group(1))
                sub_num = int(subsection_match.group(2))
                original_title = subsection_match.group(3).strip()
                
                # Find the corresponding main section
                for section in sections_found:
                    if section['original_number'] == main_num:
                        section['subsections'].append({
                            'sub_number': sub_num,
                            'title': original_title
                        })
                        break
        
        # Second pass: generate TOC with proper numbering
        for section in sections_found:
            original_title = section['title']
            
            # Map to Indonesian if available
            mapped_title = section_mappings.get(original_title.upper(), original_title)
            
            toc_items.append({
                'number': str(section_counter),
                'title': mapped_title,
                'page': page_number
            })
            
            current_main_section = section_counter
            page_number += 1
            
            # Add subsections
            for i, subsection in enumerate(section['subsections'], 1):
                sub_title = subsection['title']
                # Map subsection to Indonesian if available
                mapped_sub_title = section_mappings.get(sub_title, sub_title)
                
                toc_items.append({
                    'number': f'{current_main_section}.{i}',
                    'title': mapped_sub_title,
                    'page': page_number
                })
                page_number += 1
            
            section_counter += 1
        
        return toc_items
    
    def create_valid_datasets_table_data(self, valid_datasets: list) -> tuple:
        if not valid_datasets:
            return [], []
        headers = ['Data', 'Kondisi']
        data = [[row['data'], row['condition']] for row in valid_datasets]
        return data, headers

    def create_country_info_table_data(self, country_info: list) -> tuple:
        if not country_info:
            return [], []
        headers = ['Data', 'Kondisi']
        data = [[row['data'], row['condition']] for row in country_info]
        return data, headers


class WordTableManager:
    """Enhanced Word table manager for creating proper tables with borders"""
    
    @staticmethod
    def create_bordered_table(doc, data, headers, table_style='Table Grid'):
        """Create a properly formatted table with borders and styling"""
        from docx.shared import Inches, Pt
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.shared import OxmlElement, qn
        
        if not data or not headers:
            return None
        
        # Create table
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = table_style
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        # Set column widths based on content
        col_count = len(headers)
        if col_count == 3:  # Detail Processing table or Error handling 3-column format
            if 'Nama Field' in headers[0]:  # Detail Processing table
                table.columns[0].width = Inches(2.0)  # Nama Field
                table.columns[1].width = Inches(1.5)  # Technical Field  
                table.columns[2].width = Inches(4.0)  # Keterangan
            else:  # Generic 3-column table
                table.columns[0].width = Inches(1.5)
                table.columns[1].width = Inches(3.0)
                table.columns[2].width = Inches(3.0)
        elif col_count == 6:  # Selection Screen table
            table.columns[0].width = Inches(1.2)  # Parameter
            table.columns[1].width = Inches(1.0)  # Type
            table.columns[2].width = Inches(2.0)  # Description
            table.columns[3].width = Inches(0.8)  # Mandatory
            table.columns[4].width = Inches(1.0)  # Select-Option
            table.columns[5].width = Inches(0.8)  # No Intervals
        elif col_count == 2:  # Testing Requirements table (KONDISI PENGUJIAN | HASIL YANG DIHARAPKAN)
            table.columns[0].width = Inches(3.5)  # Kondisi Pengujian
            table.columns[1].width = Inches(4.0)  # Hasil yang Diharapkan
        elif col_count == 4:  # Error Handling table (No | Error | Resolution | Code & Severity)
            table.columns[0].width = Inches(0.5)  # No
            table.columns[1].width = Inches(3.0)  # Error Description
            table.columns[2].width = Inches(3.0)  # Resolution
            table.columns[3].width = Inches(1.0)  # Error Code & Severity
        elif col_count == 5:  # Full Testing Requirements table or Error Handling table
            if 'KONDISI PENGUJIAN' in headers[1]:  # Testing Requirements table
                table.columns[0].width = Inches(0.4)  # No
                table.columns[1].width = Inches(2.8)  # Kondisi Pengujian
                table.columns[2].width = Inches(2.8)  # Hasil yang Diharapkan
                table.columns[3].width = Inches(1.5)  # Data Uji
                table.columns[4].width = Inches(0.8)  # Prioritas
            else:  # Error Handling table
                table.columns[0].width = Inches(0.5)  # No
                table.columns[1].width = Inches(2.5)  # Error Description
                table.columns[2].width = Inches(2.5)  # Resolution
                table.columns[3].width = Inches(1.0)  # Error Code
                table.columns[4].width = Inches(0.8)  # Severity
        
        # Add and format headers
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            
            # Format header cell
            for paragraph in header_cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
                    run.font.name = 'Arial'
        
        # Add data rows
        for row_data in data:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                if i < len(row_cells):
                    row_cells[i].text = str(cell_data) if cell_data else ''
                    
                    # Format data cell
                    for paragraph in row_cells[i].paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        for run in paragraph.runs:
                            run.font.size = Pt(9)
                            run.font.name = 'Arial'
        
        # Apply additional table formatting
        WordTableManager._apply_table_borders(table)
        
        return table
    
    @staticmethod
    def _apply_table_borders(table):
        """Apply consistent borders to all table cells"""
        from docx.oxml.shared import OxmlElement, qn
        from docx.oxml.ns import nsdecls
        
        def set_cell_border(cell, **kwargs):
            """Set border for a cell"""
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            
            # Create borders element
            tcBorders = OxmlElement('w:tcBorders')
            
            for edge in ('top', 'left', 'bottom', 'right'):
                edge_data = kwargs.get(edge)
                if edge_data:
                    tag = f'w:{edge}'
                    element = OxmlElement(tag)
                    element.set(qn('w:val'), edge_data['val'])
                    element.set(qn('w:sz'), str(edge_data['sz']))
                    element.set(qn('w:color'), edge_data['color'])
                    tcBorders.append(element)
            
            tcPr.append(tcBorders)
        
        # Border settings
        border_settings = {
            'top': {'val': 'single', 'sz': '4', 'color': '000000'},
            'left': {'val': 'single', 'sz': '4', 'color': '000000'},
            'bottom': {'val': 'single', 'sz': '4', 'color': '000000'},
            'right': {'val': 'single', 'sz': '4', 'color': '000000'}
        }
        
        # Apply borders to all cells
        for row in table.rows:
            for cell in row.cells:
                set_cell_border(cell, **border_settings)

class TitlePageWordGenerator:
    """Generate Word document for title page and document information with proper tables"""
    
    def __init__(self, template_path: str = None):
        self.template_path = template_path
    
    def create_selection_screen_table_data(self, selection_screen_table: list) -> tuple:
        """Prepare data for Selection Screen table"""
        if not selection_screen_table:
            return [], []
        
        headers = ['Parameter', 'Type', 'Description', 'Mandatory', 'Select-Option', 'No Intervals']
        data = []
        
        for item in selection_screen_table:
            row = [
                item.get('parameter', 'N/A'),
                item.get('type', 'N/A'),
                item.get('description', 'N/A'),
                item.get('mandatory', 'No'),
                item.get('select_option', 'No'),
                item.get('no_intervals', 'No')
            ]
            data.append(row)
        
        return data, headers
    
    def create_detail_processing_table_data(self, detail_processing_table: list) -> tuple:
        """Prepare data for Detail Processing table"""
        if not detail_processing_table:
            return [], []
        
        headers = ['Nama Field', 'Technical Field', 'Keterangan']
        data = []
        
        for item in detail_processing_table:
            # Create comprehensive Keterangan
            keterangan_parts = []
            
            # Add source table info
            if item.get('source_table'):
                keterangan_parts.append(f"Ambil {item.get('technical_field', 'N/A')} dari tabel {item.get('source_table')}")
            
            # Add processing logic
            if item.get('processing_logic') and item.get('processing_logic').strip() != 'N/A':
                logic = item.get('processing_logic', '').strip()
                if logic:
                    keterangan_parts.append(logic)
            
            # Add processing type if different from DIRECT
            if item.get('processing_type') and item.get('processing_type') != 'DIRECT':
                keterangan_parts.append(f"Processing Type: {item.get('processing_type')}")
            
            keterangan = '. '.join(keterangan_parts) if keterangan_parts else 'N/A'
            
            row = [
                item.get('field_name', 'N/A'),
                item.get('technical_field', 'N/A'),
                keterangan
            ]
            data.append(row)
        
        return data, headers

    # ───── Add these here ───────────────────────────────────────────────────
    def create_valid_datasets_table_data(self, valid_datasets: list) -> tuple:
        if not valid_datasets:
            return [], []
        headers = ['Data', 'Kondisi']
        data = [[row['data'], row['condition']] for row in valid_datasets]
        return data, headers

    def create_country_info_table_data(self, country_info: list) -> tuple:
        if not country_info:
            return [], []
        headers = ['Data', 'Kondisi']
        data = [[row['data'], row['condition']] for row in country_info]
        return data, headers
    
    def create_error_handling_table_data(self, error_handling_table: list) -> tuple:
        """Prepare data for Error Handling table"""
        if not error_handling_table:
            return [], []
        
        # Use 4-column format: No | Error Description | Resolution | Error Code & Severity
        headers = ['No', 'Potensi Error', 'Penyelesaian', 'Kode Error & Severity']
        data = []
        
        for item in error_handling_table:
            error_desc = item.get('error_description', 'N/A')
            resolution = item.get('resolution', 'N/A')
            error_code = item.get('error_code', 'N/A')
            severity = item.get('severity', 'ERROR')
            
            # Combine error code and severity
            code_severity = f"{error_code} ({severity})" if error_code != 'N/A' else f"({severity})"
            
            row = [
                item.get('no', str(len(data) + 1)),
                error_desc,
                resolution,
                code_severity
            ]
            data.append(row)
        
        return data, headers
    
    def create_testing_requirements_table_data(self, testing_requirements_table: list) -> tuple:
        """Prepare data for Testing Requirements table - CLEAN VERSION"""
        if not testing_requirements_table:
            return [], []
        
        # Use the full 5-column format from markdown: No | Test Condition | Expected Result | Test Data | Priority
        headers = ['No', 'KONDISI PENGUJIAN', 'HASIL YANG DIHARAPKAN', 'DATA UJI', 'PRIORITAS']
        data = []
        
        for item in testing_requirements_table:
            row = [
                item.get('no', str(len(data) + 1)),
                item.get('test_condition', 'N/A'),
                item.get('expected_result', 'N/A'),
                item.get('test_data', 'N/A'),
                item.get('priority', 'MEDIUM').upper()
            ]
            data.append(row)
        
        return data, headers
    
    def create_testing_requirements_summary_table_data(self, testing_requirements_table: list) -> tuple:
        """Create simplified 2-column table for testing requirements if space is limited"""
        if not testing_requirements_table:
            return [], []
        
        headers = ['KONDISI PENGUJIAN', 'HASIL YANG DIHARAPKAN']
        data = []
        
        for item in testing_requirements_table:
            row = [
                item.get('test_condition', 'N/A'),
                item.get('expected_result', 'N/A')
            ]
            data.append(row)
        
        return data, headers
    
    def create_data_condition_table_data(self, table_rows: list) -> tuple:
        """Create simple Data/Kondisi table"""
        if not table_rows:
            return [], []

        headers = ['Data', 'Kondisi']
        data = []
        for item in table_rows:
            row = [item.get('data', 'N/A'), item.get('condition', 'N/A')]
            data.append(row)

        return data, headers
    
    def generate_testing_requirements_content_text(self, testing_requirements_table: list) -> str:
        """Generate testing requirements content as text summary - SIMPLIFIED VERSION"""
        if not testing_requirements_table:
            return "Melakukan pengujian menyeluruh terhadap semua fungsi dan fitur program untuk memastikan kesesuaian dengan spesifikasi yang telah ditentukan."
        
        # SIMPLIFIED - just return basic intro text without the detailed breakdown
        return "Pengujian harus dilakukan pada environment development dengan data representatif yang mencakup berbagai skenario bisnis."
    
    def replace_text_preserve_formatting(self, paragraph, old_text: str, new_text: str):
        """Replace text while preserving formatting"""
        if old_text in paragraph.text:
            # Find all runs and their text
            full_text = paragraph.text
            if old_text in full_text:
                # Get the replacement position
                new_full_text = full_text.replace(old_text, new_text)
                
                # Clear existing runs but keep formatting from first run
                if paragraph.runs:
                    first_run = paragraph.runs[0]
                    # Preserve formatting properties
                    font_name = first_run.font.name
                    font_size = first_run.font.size
                    font_bold = first_run.font.bold
                    font_italic = first_run.font.italic
                    font_color = first_run.font.color.rgb if first_run.font.color.rgb else None
                    
                    # Clear all runs
                    paragraph.clear()
                    
                    # Handle multiline text (for assumptions)
                    if '\n' in new_full_text:
                        lines = new_full_text.split('\n')
                        for i, line in enumerate(lines):
                            if i > 0:
                                # Add line break
                                line_break_run = paragraph.add_run()
                                line_break_run.add_break()
                            new_run = paragraph.add_run(line)
                            if font_name:
                                new_run.font.name = font_name
                            if font_size:
                                new_run.font.size = font_size
                            if font_bold:
                                new_run.font.bold = font_bold
                            if font_italic:
                                new_run.font.italic = font_italic
                            if font_color:
                                new_run.font.color.rgb = font_color
                    else:
                        # Single line text
                        new_run = paragraph.add_run(new_full_text)
                        if font_name:
                            new_run.font.name = font_name
                        if font_size:
                            new_run.font.size = font_size
                        if font_bold:
                            new_run.font.bold = font_bold
                        if font_italic:
                            new_run.font.italic = font_italic
                        if font_color:
                            new_run.font.color.rgb = font_color
                else:
                    # If no runs, just set paragraph text
                    paragraph.text = new_full_text
                
                logger.info(f"Replaced '{old_text[:30]}...' with '{new_text[:30]}...' (formatting preserved)")
                return True
        return False
    
    def find_and_insert_table_after_text(self, doc, search_text: str, table_data: list, headers: list, title: str = None) -> bool:
        """Find text and insert table after it"""
        try:
            inserted = False
            
            # Search through all paragraphs
            for i, paragraph in enumerate(doc.paragraphs):
                if search_text in paragraph.text:
                    logger.info(f"Found search text: '{search_text}' in paragraph {i}")
                    
                    # Clear the paragraph or replace it with title
                    paragraph.clear()
                    if title:
                        title_run = paragraph.add_run(title)
                        title_run.font.bold = True
                        title_run.font.size = Pt(12)
                        title_run.font.name = 'Arial'
                    
                    # Create the table
                    table = WordTableManager.create_bordered_table(doc, table_data, headers)
                    if table:
                        # Insert table after the paragraph
                        paragraph._element.addnext(table._element)
                        logger.info(f"✅ Inserted table with {len(table_data)} rows after '{search_text}'")
                        inserted = True
                        break
            
            # If not found in paragraphs, search in table cells
            if not inserted:
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                if search_text in paragraph.text:
                                    # Replace with a summary
                                    paragraph.clear()
                                    summary_text = f"{title}: {len(table_data)} items" if title else f"Table with {len(table_data)} rows"
                                    paragraph.text = summary_text
                                    logger.info(f"✅ Replaced placeholder in table cell: '{search_text}'")
                                    inserted = True
                                    break
                            if inserted:
                                break
                        if inserted:
                            break
                    if inserted:
                        break
            
            return inserted
            
        except Exception as e:
            logger.error(f"❌ Error inserting table after '{search_text}': {e}")
            return False
    
    def find_and_replace_testing_section(self, doc, testing_data: list, testing_headers: list) -> bool:
        """Specifically handle testing requirements section - CLEAN VERSION WITHOUT REDUNDANT TEXT"""
        try:
            inserted = False
            
            # Look for multiple possible locations for testing requirements
            search_patterns = [
                '{{persyaratan_pengujian}}',
                'Persyaratan pengujian yang harus dipenuhi:',
                'Prioritas Tinggi:',
                '[*Masukkan kondisi fungsional yang diperlukan untuk pengujian.*]',
                '**KONDISI PENGUJIAN**',
                'KONDISI PENGUJIAN',
                'Persyaratan pengujian mencakup skenario'  # New pattern to catch the generated summary
            ]
            
            for pattern in search_patterns:
                for i, paragraph in enumerate(doc.paragraphs):
                    if pattern in paragraph.text and not inserted:
                        logger.info(f"Found testing pattern: '{pattern}' in paragraph {i}")
                        
                        # Clear this paragraph completely - NO INTRO TEXT
                        paragraph.clear()
                        
                        # Create the table directly without any introduction
                        table = WordTableManager.create_bordered_table(doc, testing_data, testing_headers)
                        if table:
                            # Insert table directly
                            paragraph._element.addnext(table._element)
                            logger.info(f"✅ Inserted testing requirements table with {len(testing_data)} rows")
                            inserted = True
                            
                            # Clean up any remaining testing-related paragraphs - ENHANCED CLEANUP
                            self._cleanup_testing_paragraphs_aggressive(doc, i)
                            break
                
                if inserted:
                    break
            
            return inserted
            
        except Exception as e:
            logger.error(f"❌ Error replacing testing section: {e}")
            return False
    
    def _cleanup_testing_paragraphs_aggressive(self, doc, start_index: int):
        """Aggressively clean up ALL testing-related paragraphs that might contain redundant content"""
        cleanup_patterns = [
            'Prioritas Tinggi:',
            'Prioritas Menengah:',
            'Prioritas Rendah:',
            'User has authorization',
            'User does not have authorization',
            'Process RHR006_01',
            'Selection option s_lgart',
            'Infotype 0015',
            'Large dataset',
            'Employee has IT0015 record',
            '[*Masukkan kondisi fungsional',
            '**KONDISI PENGUJIAN**',
            '**HASIL YANG DIHARAPKAN**',
            'Persyaratan pengujian mencakup skenario',
            'skenario prioritas tinggi',
            'skenario prioritas menengah',
            'skenario prioritas rendah',
            'Detail kondisi pengujian tercantum',
            'authorization, security',
            'functional testing',
            'performance, edge cases'
        ]
        
        paragraphs_to_clean = []
        
        # Look for paragraphs that contain ANY of the cleanup patterns - EXPAND SEARCH RANGE
        for i in range(max(0, start_index - 5), min(start_index + 30, len(doc.paragraphs))):
            if i < len(doc.paragraphs):
                paragraph = doc.paragraphs[i]
                paragraph_text = paragraph.text.strip()
                
                # Check if paragraph contains any cleanup pattern
                for pattern in cleanup_patterns:
                    if pattern in paragraph_text:
                        paragraphs_to_clean.append(paragraph)
                        logger.info(f"Marked for cleanup: '{paragraph_text[:50]}...'")
                        break
                
                # Also clean paragraphs that are mostly just priority breakdowns
                if len(paragraph_text) > 50 and any(word in paragraph_text.lower() for word in ['prioritas', 'skenario', 'authorization', 'functional']):
                    if paragraph not in paragraphs_to_clean:
                        paragraphs_to_clean.append(paragraph)
                        logger.info(f"Marked priority content for cleanup: '{paragraph_text[:50]}...'")
        
        # Clean the paragraphs
        for paragraph in paragraphs_to_clean:
            original_text = paragraph.text[:50] if paragraph.text else "empty"
            paragraph.clear()
            logger.info(f"✅ Cleaned up testing paragraph: '{original_text}...'")
    
    def generate_with_proper_tables(self, title_info: dict, output_path: str) -> bool:
        """Generate Word document with proper tables replacing placeholders - CLEAN VERSION"""
        if not self.template_path or not os.path.exists(self.template_path):
            logger.error("Template path not found")
            return False
        
        try:
            # Open the template document
            doc = DocxDocument(self.template_path)
            logger.info(f"📄 Opened template: {self.template_path}")

            print("[Debug] title_info", title_info)
            ricefw   = title_info.get('ricefw_id', '').strip()      # e.g. "RHR006"
            filename = title_info.get('file_name', '').strip()      # e.g. "Laporan tunjangan & potongan tidak tetap"
            new_heading = f"{ricefw} {filename}"
            print("[Debug] New heading:", new_heading)
            for paragraph in doc.paragraphs:
                print("[Debug] paragraph", paragraph.text)
                if paragraph.text.strip().endswith("(Nama File)"):
                    print("[Debug] Found placeholder for title:", paragraph.text.strip().endswith("(Nama File)"))
                    # replace entire placeholder line with our dynamic heading
                    self.replace_text_preserve_formatting(
                        paragraph,
                        paragraph.text.strip(),
                        new_heading
                    )
                    logger.info(f"Replaced title placeholder with dynamic heading: '{new_heading}'")
                    break
            
            # Prepare content for replacements
            assumptions_text = ""
            if title_info.get('assumptions'):
                assumptions_list = []
                for i, assumption in enumerate(title_info['assumptions'], 1):
                    assumptions_list.append(f"({i}) {assumption}")
                assumptions_text = "\n".join(assumptions_list)
            else:
                assumptions_text = "(1) Tidak ada asumsi khusus"
            
            transaction_menu_text = ""
            transaction_code = title_info.get('transaction_code', '').strip()
            menu_path = title_info.get('menu_path', '').strip()
            
            if transaction_code and transaction_code != 'N/A':
                transaction_menu_text = f"Transaksi: {transaction_code}"
                if menu_path and menu_path != 'N/A':
                    transaction_menu_text += f"\nMenu Path: {menu_path}"
            elif menu_path and menu_path != 'N/A':
                transaction_menu_text = f"Menu Path: {menu_path}"
            else:
                transaction_menu_text = "N/A"
            
            # Generate DESAIN section content
            report_description = self._generate_report_description(title_info)
            authorization_info = self._generate_authorization_info(title_info)
            design_constraints = self._generate_design_constraints(title_info)
            
            # Generate CLEAN testing requirements content - SIMPLIFIED VERSION
            testing_requirements_content = self.generate_testing_requirements_content_text(title_info.get('testing_requirements_table', []))
            
            # Define standard text replacements
            replacements = {
                # Existing replacements
                '(Nama Modul)': '',
                'Human Resource (Nama Modul)': title_info.get('module_name', 'Human Resource'),
                '(DAPI ID)': title_info.get('ricefw_id', 'N/A'),
                'Functional Specification Design (FSD)': f"{title_info.get('document_type', 'Functional Specification Design (FSD)')}",
                f"{title_info.get('ricefw_id', 'RHR041')} (DAPI ID) Functional Specification Design (FSD)":
                title_info.get('document_type', 'Functional Specification Design (FSD)'),
                '(Nama File)': '',
                # 'RHR041 Travel Requisition Process (Nama File)': f"{title_info.get('ricefw_id', 'RHR041')} {title_info.get('file_name', 'Process')}",
                # '(DAPI ID) Travel Requisition Process': title_info.get('document_type', 'Travel Requisition Process'),
                '(Link NAS)': title_info.get('document_location', 'N/A'),
                'Draft awal oleh': 'AI Generated',
                'System Integrator for Management Information System Towards Single Source of Truth Implementation Program': title_info.get('project_name', 'System Integrator for Management Information System Towards Single Source of Truth Implementation Program'),
                
                # Page 4 content replacements
                '{{persyaratan_pengguna}}': title_info.get('user_requirements', 'Tidak ada persyaratan khusus yang didefinisikan.'),
                '{{asumsi}}': assumptions_text,
                '{{nama_program_sap}}': title_info.get('sap_program_name', 'N/A'),
                '{{transaksi_menu_sap}}': transaction_menu_text,
                
                # DESAIN section replacements
                '{{deskripsi_detail_report}}': report_description,
                '{{otorisasi}}': authorization_info,
                '{{keterbatasan_desain}}': design_constraints,
                
                # DESIGN ALTERNATIVES section replacements
                '{{latar_belakang}}': 'N/A',
                '{{opsi}}': 'N/A',
                '{{rekomendasi}}': 'N/A',
                
                # TESTING REQUIREMENTS section replacements - CLEAN VERSION
                '{{persyaratan_pengujian}}': testing_requirements_content,
                '[*Masukkan kondisi fungsional yang diperlukan untuk pengujian.*]': '',  # Remove this entirely
                '{{data_uji}}': title_info.get('test_data_location', 'Data uji tersedia di environment development SAP'),
                '{{transaksi}}': title_info.get('test_transaction', 'N/A'),
                '{{menu_path}}': title_info.get('test_menu_path', 'N/A')
            }

            file_pattern = re.compile(r"RHR\d+[^\n]*\(Nama File\)")
            doc_pattern = re.compile(r"RHR\d+\s*\(DAPI ID\)\s*Functional Specification Design \(FSD\)")
            
            # Replace standard text in all paragraphs
            for paragraph in doc.paragraphs:
                for old_text, new_text in replacements.items():
                    if old_text in paragraph.text:
                        self.replace_text_preserve_formatting(paragraph, old_text, new_text)
            
            match_file = file_pattern.search(paragraph.text)
            if match_file:
                self.replace_text_preserve_formatting(
                    paragraph,
                    match_file.group(),
                    f"{title_info.get('ricefw_id', 'RHR041')} {title_info.get('file_name', 'Process')}"
                )

            match_doc = doc_pattern.search(paragraph.text)
            if match_doc:
                self.replace_text_preserve_formatting(
                    paragraph,
                    match_doc.group(),
                    title_info.get('document_type', 'Functional Specification Design (FSD)')
                )
            # Replace text in table cells
            for table in doc.tables:
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        for paragraph in cell.paragraphs:
                            # Apply standard replacements
                            for old_text, new_text in replacements.items():
                                if old_text in paragraph.text:
                                    self.replace_text_preserve_formatting(paragraph, old_text, new_text)
                            
                            # Handle version history date
                            if ("AI Generated" in paragraph.text or "0.01" in paragraph.text) and paragraph.text.strip() == "":
                                row_texts = [c.text.strip() for c in row.cells]
                                if any("0.01" in text or "AI Generated" in text for text in row_texts):
                                    if cell_idx == len(row.cells) - 1:
                                        paragraph.text = title_info.get('generated_date', '')
                            
                            if paragraph.text.strip() == "":
                                row_texts = [c.text.strip() for c in row.cells]
                                if "AI Generated" in " ".join(row_texts) and cell_idx == len(row.cells) - 1:
                                    paragraph.text = title_info.get('generated_date', '')
            
            # CREATE AND INSERT TABLES - CLEAN VERSION
            
            # 1. Handle Selection Screen table
            selection_data, selection_headers = self.create_selection_screen_table_data(title_info.get('selection_screen_table', []))
            if selection_data and selection_headers:
                logger.info(f"🔧 Creating Selection Screen table with {len(selection_data)} rows")
                self.find_and_insert_table_after_text(
                    doc, 
                    'Selection Screen', 
                    selection_data, 
                    selection_headers, 
                    "Selection Screen:"
                )
            
            # 2. Handle Detail Processing table
            detail_data, detail_headers = self.create_detail_processing_table_data(title_info.get('detail_processing_table', []))
            if detail_data and detail_headers:
                logger.info(f"🔧 Creating Detail Processing table with {len(detail_data)} rows")
                self.find_and_insert_table_after_text(
                    doc,
                    'Detail Processing',
                    detail_data,
                    detail_headers,
                    "Detail Processing:"
                )

            # 2b. Handle additional design tables
            valid_data, valid_headers = self.create_data_condition_table_data(title_info.get('valid_datasets_table', []))
            if valid_data:
                logger.info(f"🔧 Creating Valid Dataset table with {len(valid_data)} rows")
                self.find_and_insert_table_after_text(
                    doc,
                    'Detail Process Only valid datasets',
                    valid_data,
                    valid_headers,
                    "Detail Process Only valid datasets:"
                )

            country_data, country_headers = self.create_data_condition_table_data(title_info.get('country_info_table', []))
            if country_data:
                logger.info(f"🔧 Creating Country Info table with {len(country_data)} rows")
                self.find_and_insert_table_after_text(
                    doc,
                    'Form Get_Country_Info',
                    country_data,
                    country_headers,
                    "Form Get_Country_Info:"
                )

            c500c_data, c500c_headers = self.create_data_condition_table_data(title_info.get('currency_t500c_table', []))
            if c500c_data:
                logger.info(f"🔧 Creating Currency T500C table with {len(c500c_data)} rows")
                self.find_and_insert_table_after_text(
                    doc,
                    'Form Get_Currency_T500C',
                    c500c_data,
                    c500c_headers,
                    "Form Get_Currency_T500C:"
                )

            c001_data, c001_headers = self.create_data_condition_table_data(title_info.get('currency_t001_table', []))
            if c001_data:
                logger.info(f"🔧 Creating Currency T001 table with {len(c001_data)} rows")
                self.find_and_insert_table_after_text(
                    doc,
                    'Form Get_Currency_T001',
                    c001_data,
                    c001_headers,
                    "Form Get_Currency_T001:"
                )
            
            # 4. Handle Error Handling table
            error_data, error_headers = self.create_error_handling_table_data(title_info.get('error_handling_table', []))
            if error_data and error_headers:
                logger.info(f"🔧 Creating Error Handling table with {len(error_data)} rows")
                
                # Replace the error handling content with proper table
                self.find_and_insert_table_after_text(
                    doc, 
                    '{{potensi_error}}', 
                    error_data, 
                    error_headers, 
                    "Potensi Error:"
                )
                
                # Also replace any placeholder for notification procedures
                for paragraph in doc.paragraphs:
                    if '[*Masukkan deskripsi tentang error log, laporan, dan/atau pesan yang relevan.*]' in paragraph.text:
                        self.replace_text_preserve_formatting(
                            paragraph, 
                            '[*Masukkan deskripsi tentang error log, laporan, dan/atau pesan yang relevan.*]',
                            'Error akan dicatat dalam log sistem SAP dan dapat dilihat melalui transaction ST22 (Dump Analysis) atau SLG1 (Application Log).'
                        )
            
            # 5. Handle Testing Requirements table - CLEAN VERSION WITH NO REDUNDANT TEXT
            testing_data, testing_headers = self.create_testing_requirements_table_data(title_info.get('testing_requirements_table', []))
            if testing_data and testing_headers:
                logger.info(f"🔧 Creating Testing Requirements table with {len(testing_data)} rows - CLEAN VERSION")
                
                # Use the specialized testing section replacement method that removes redundant text
                testing_inserted = self.find_and_replace_testing_section(doc, testing_data, testing_headers)
                
                if not testing_inserted:
                    # Fallback: try the regular method
                    self.find_and_insert_table_after_text(
                        doc, 
                        '{{kondisi_pengujian_tabel}}', 
                        testing_data, 
                        testing_headers, 
                        None  # NO TITLE to avoid redundancy
                    )
            
            # ENHANCED Final cleanup: remove any remaining placeholders and redundant text
            placeholders_to_remove = [
                '{{detail_processing}}',
                '{{kondisi_pengujian_tabel}}',
                '{{potensi_error}}',
                'Prioritas Tinggi: - User has authorization',
                'Prioritas Menengah: - Process RHR006',
                'Prioritas Rendah: - Employee has IT0015',
                '**KONDISI PENGUJIAN**',
                '**HASIL YANG DIHARAPKAN**',
                'Persyaratan pengujian mencakup skenario berikut:',
                'Detail kondisi pengujian tercantum dalam tabel di bawah ini.',
                'skenario prioritas tinggi (authorization, security)',
                'skenario prioritas menengah (functional testing)',
                'skenario prioritas rendah (performance, edge cases)',
                '[*Masukkan kondisi fungsional yang diperlukan untuk pengujian.*]'
            ]
            
            for paragraph in doc.paragraphs:
                original_text = paragraph.text
                should_clear = False
                
                for placeholder in placeholders_to_remove:
                    if placeholder in paragraph.text:
                        # If the paragraph is mostly the placeholder, clear it entirely
                        if len(paragraph.text.strip()) <= len(placeholder) + 50:
                            should_clear = True
                            break
                        else:
                            # Remove the placeholder but keep other text
                            paragraph.text = paragraph.text.replace(placeholder, '').strip()
                
                # Clear paragraphs that are mostly redundant testing content
                if not should_clear and paragraph.text.strip():
                    text_lower = paragraph.text.lower()
                    if (len(paragraph.text.strip()) < 200 and 
                        any(phrase in text_lower for phrase in ['prioritas tinggi', 'prioritas menengah', 'prioritas rendah', 'skenario']) and
                        any(phrase in text_lower for phrase in ['authorization', 'functional', 'performance'])):
                        should_clear = True
                
                if should_clear:
                    logger.info(f"Clearing redundant paragraph: '{original_text[:50]}...'")
                    paragraph.clear()
            
            # Save the document
            doc.save(output_path)
            logger.info(f"✅ Successfully generated Word document with CLEAN tables (no redundant text): {output_path}")
            
            # Verify file was created and has content
            if os.path.exists(output_path):
                file_size = os.path.getsize(output_path)
                logger.info(f"📏 Generated file size: {file_size:,} bytes")
                return True
            else:
                logger.error("❌ File was not created")
                return False
            
        except Exception as e:
            logger.error(f"❌ Error generating Word document with tables: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    def _generate_report_description(self, title_info: dict) -> str:
        """Generate report description based on extracted information"""
        description_parts = []
        
        # Use user requirements as base description
        if title_info.get('user_requirements'):
            description_parts.append(title_info['user_requirements'])
        
        # Add information about selection screen if available
        if title_info.get('selection_screen_table'):
            param_count = len(title_info['selection_screen_table'])
            description_parts.append(f"Report ini memiliki {param_count} parameter selection screen.")
        
        # Add information about processing fields if available
        if title_info.get('detail_processing_table'):
            field_count = len(title_info['detail_processing_table'])
            description_parts.append(f"Proses detail melibatkan {field_count} field utama.")
        
        if not description_parts:
            description_parts.append("Report ini dirancang untuk memenuhi kebutuhan bisnis sesuai spesifikasi yang telah ditentukan.")
        
        return " ".join(description_parts)
    
    def _generate_authorization_info(self, title_info: dict) -> str:
        """Generate authorization information"""
        auth_parts = []
        
        # Check if there are any authorization-related assumptions
        if title_info.get('assumptions'):
            for assumption in title_info['assumptions']:
                if 'authorization' in assumption.lower() or 'authorizations' in assumption.lower():
                    auth_parts.append(assumption)
        
        if not auth_parts:
            auth_parts.append("User harus memiliki otorisasi yang sesuai untuk mengakses dan menjalankan report ini.")
        
        return "\n".join(auth_parts)
    
    def _generate_design_constraints(self, title_info: dict) -> str:
        """Generate design constraints information"""
        constraints_parts = []
        
        # Add constraints based on technical requirements
        if title_info.get('sap_program_name'):
            constraints_parts.append(f"Report harus kompatibel dengan program SAP {title_info['sap_program_name']}.")
        
        # Add constraints from assumptions that might be technical constraints
        if title_info.get('assumptions'):
            for assumption in title_info['assumptions']:
                if any(keyword in assumption.lower() for keyword in ['locking', 'performance', 'data', 'infotype']):
                    constraints_parts.append(assumption)
        
        if not constraints_parts:
            constraints_parts.append("Report harus mengikuti standard development guidelines dan best practices SAP.")
        
        return "\n".join(constraints_parts)

class TitlePageGenerator:
    """Main class to generate title page and document information documents with proper tables"""
    
    def __init__(self, template_path: str = None):
        self.template_path = template_path
        self.word_generator = TitlePageWordGenerator(template_path)
    
    def generate_complete_document(self, markdown_path: str, output_dir: str) -> dict:
        """Generate complete documents with proper Word tables from markdown"""
        try:
            # Create output directory
            Path(output_dir).mkdir(parents=True, exist_ok=True)
            
            # Read markdown file
            logger.info(f"📖 Reading markdown file: {markdown_path}")
            with open(markdown_path, 'r', encoding='utf-8') as f:
                markdown_content = f.read()
            
            # Extract title information
            logger.info("🔍 Extracting title and document information...")
            extractor = MarkdownTitleExtractor(markdown_content)
            title_info = extractor.extract_title_info()
            
            # Generate timestamp for unique filenames
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            base_filename = f"FSD_Complete_{title_info.get('program_name', 'Document')}_{timestamp}"
            
            # Generate Word document with proper tables
            logger.info("📝 Generating Word document with CLEAN bordered tables...")
            word_path = os.path.join(output_dir, f"{base_filename}.docx")
            
            # Generate using the enhanced table method
            word_success = self.word_generator.generate_with_proper_tables(title_info, word_path)
            
            # Prepare results
            results = {
                'word_path': word_path if word_success else None,
                'title_info': title_info,
                'success': word_success,
                'tables_created': {
                    'selection_screen': len(title_info.get('selection_screen_table', [])),
                    'detail_processing': len(title_info.get('detail_processing_table', [])),
                    'error_handling': len(title_info.get('error_handling_table', [])),
                    'testing_requirements': len(title_info.get('testing_requirements_table', []))
                }
            }
            
            return results
            
        except Exception as e:
            logger.error(f"❌ Error generating complete document: {e}")
            raise

def main():
    """Main execution function for complete document generation with CLEAN tables"""
    
    # Configuration - DO NOT HARDCODE VALUES, USE PARAMETERS OR CONFIG FILES
    import sys
    
    # Default values - these should be passed as parameters or from config
    # DEFAULT_MARKDOWN_PATH = "/Users/wilbert.limson/python_project/PLN-Genie/output/zhr_r_it0267_fsd.md"
    DEFAULT_MARKDOWN_PATH = "/Users/wilbert.limson/python_project/PLN-Genie/output/zhr_r_it0015_fsd.md"
    DEFAULT_TEMPLATE_PATH = "/Users/wilbert.limson/python_project/PLN-Genie/Template/FSD/Template_PLN_SI SSoT_(DAPI ID)_(Module Name)_Functional Specification Design (FSD)_v100_ID.docx"
    DEFAULT_OUTPUT_DIR = "/Users/wilbert.limson/python_project/PLN-Genie/one_output/"
    
    # Allow command line arguments
    if len(sys.argv) >= 4:
        MARKDOWN_PATH = sys.argv[1]
        TEMPLATE_PATH = sys.argv[2] 
        OUTPUT_DIR = sys.argv[3]
    else:
        MARKDOWN_PATH = DEFAULT_MARKDOWN_PATH
        TEMPLATE_PATH = DEFAULT_TEMPLATE_PATH
        OUTPUT_DIR = DEFAULT_OUTPUT_DIR
    
    try:
        # Validate input files
        if not os.path.exists(MARKDOWN_PATH):
            raise FileNotFoundError(f"Markdown file not found: {MARKDOWN_PATH}")
        
        if not os.path.exists(TEMPLATE_PATH):
            logger.warning(f"Template file not found: {TEMPLATE_PATH}")
            logger.info("Will generate documents without Word template")
            TEMPLATE_PATH = None
        
        # Initialize title page generator
        logger.info("🚀 Starting complete document generation with CLEAN table formatting (no redundant text)...")
        generator = TitlePageGenerator(TEMPLATE_PATH)
        
        # Generate complete document
        results = generator.generate_complete_document(MARKDOWN_PATH, OUTPUT_DIR)
        
        # Print results
        print("\n✅ CLEAN Complete Document Generation with All Sections Completed!")
        print("=" * 90)
        
        title_info = results['title_info']
        print("📋 Extracted Information:")
        print(f"   Program Name: {title_info['program_name']}")
        print(f"   RICEFW ID: {title_info['ricefw_id']}")
        print(f"   Description: {title_info['description']}")
        print(f"   Selection Screen Parameters: {len(title_info['selection_screen_table'])} items")
        print(f"   Detail Processing Fields: {len(title_info['detail_processing_table'])} items")
        print(f"   Error Handling Scenarios: {len(title_info['error_handling_table'])} items")
        print(f"   Testing Requirements: {len(title_info['testing_requirements_table'])} items")
        
        if title_info.get('error_handling_table'):
            print(f"\n🚨 Error Handling Scenarios:")
            for i, error in enumerate(title_info['error_handling_table'], 1):
                print(f"   {i}. {error.get('error_description', 'N/A')} (Code: {error.get('error_code', 'N/A')})")
        
        if title_info.get('testing_requirements_table'):
            print(f"\n🧪 Testing Requirements:")
            high_priority = [item for item in title_info['testing_requirements_table'] if item.get('priority', '').upper() == 'HIGH']
            medium_priority = [item for item in title_info['testing_requirements_table'] if item.get('priority', '').upper() == 'MEDIUM']
            low_priority = [item for item in title_info['testing_requirements_table'] if item.get('priority', '').upper() == 'LOW']
            
            print(f"   High Priority: {len(high_priority)} tests")
            print(f"   Medium Priority: {len(medium_priority)} tests")
            print(f"   Low Priority: {len(low_priority)} tests")
        
        print("\n📁 Generated Files:")
        if results['word_path'] and results['success']:
            print(f"   📝 Word: {results['word_path']}")
            print(f"   🔧 Tables Created:")
            print(f"      📊 Selection Screen: {results['tables_created']['selection_screen']} parameters")
            print(f"      🔧 Detail Processing: {results['tables_created']['detail_processing']} fields")
            print(f"      🚨 Error Handling: {results['tables_created']['error_handling']} scenarios")
            print(f"      🧪 Testing Requirements: {results['tables_created']['testing_requirements']} test cases")
        else:
            print("   📝 Word: ❌ Failed to generate")
        
        # Check file sizes
        if results['word_path'] and os.path.exists(results['word_path']):
            word_size = os.path.getsize(results['word_path'])
            print(f"\n📏 Word file size: {word_size:,} bytes")
        
        print("=" * 90)
        
        if results['success']:
            print("🎉 SUCCESS! Complete document with CLEAN table formatting generated!")
            print("📝 FIXES Applied:")
            print("   ✅ REMOVED REDUNDANT TEXT in Testing Requirements:")
            print("      - No more 'Persyaratan pengujian mencakup skenario berikut:'")
            print("      - No more priority breakdown text before table")
            print("      - No more 'Detail kondisi pengujian tercantum dalam tabel di bawah ini.'")
            print("      - Removed placeholder text '[*Masukkan kondisi fungsional...*]'")
            print("   ✅ CLEAN Testing Requirements section:")
            print("      - Shows only the table with proper headers")
            print("      - 5-column format: No | KONDISI PENGUJIAN | HASIL YANG DIHARAPKAN | DATA UJI | PRIORITAS")
            print("      - All redundant summary text removed")
            print("   ✅ Enhanced cleanup logic:")
            print("      - Aggressive cleanup of testing-related redundant content")
            print("      - Better pattern matching for redundant text removal")
            print("      - Expanded search range for cleanup")
            print("   ✅ All other sections maintained:")
            print("      - Error Handling table still works perfectly")
            print("      - Selection Screen and Detail Processing tables unchanged")
            print("      - All table borders and formatting preserved")
        else:
            print("❌ FAILED! Document generation unsuccessful")
        
    except FileNotFoundError as e:
        print(f"❌ File not found: {e}")
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()