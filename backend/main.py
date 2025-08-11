import os
import re
import json
import logging
import asyncio
import aiohttp
import traceback
from bs4 import BeautifulSoup
from datetime import datetime
from dataclasses import dataclass, field, asdict, is_dataclass
from enum import Enum
from typing import Dict, List, Optional, Any, Union, Set
from pathlib import Path
import markdown
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx import Document as DocxDocument
from docxtpl import DocxTemplate
from md_to_docs_converter import TitlePageGenerator
import markdown2
from io import StringIO
import openpyxl
import uuid
import tempfile
import shutil
import base64

# FastAPI imports
from fastapi import FastAPI, File, UploadFile, HTTPException, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, FileResponse
from pydantic import BaseModel

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# ================================
# CORE FSD DATA STRUCTURES
# ================================

class FieldProcessingType(Enum):
    """Types of field processing in ABAP"""
    DIRECT = "DIRECT"
    LOOKUP = "LOOKUP"
    CALCULATION = "CALCULATION"
    CONSTANT = "CONSTANT"
    AGGREGATION = "AGGREGATION"

@dataclass
class FieldMapping:
    """Represents a field mapping in Detail Processing"""
    display_name: str
    technical_field: str
    source_table: str
    processing_logic: str
    processing_type: FieldProcessingType
    join_condition: str = ""
    where_condition: str = ""

@dataclass
class SelectionParameter:
    """Selection screen parameter"""
    name: str
    type: str
    description: str
    is_mandatory: bool = False
    is_select_option: bool = False
    has_no_intervals: bool = False
    default_value: str = ""

@dataclass
class ErrorScenario:
    """Error handling scenario"""
    error_description: str
    resolution: str
    error_code: str = ""
    severity: str = "ERROR"

@dataclass
class TestScenario:
    """Test scenario"""
    condition: str
    expected_result: str
    test_data: str = ""
    priority: str = "HIGH"

@dataclass
class DataConditionRow:
    """Simple table row with data and condition"""
    data: str
    condition: str

@dataclass
class FSDDocument:
    """Complete FSD Document Structure"""
    # Document Information
    project_name: str = ""
    document_location: str = ""
    related_documents: List[str] = field(default_factory=list)
    reviewers: List[Dict[str, str]] = field(default_factory=list)
    version_history: List[Dict[str, str]] = field(default_factory=list)
    
    # General Requirements
    user_requirements: str = ""
    assumptions: List[str] = field(default_factory=list)
    
    # Existing SAP Objects
    program_name: str = ""
    transaction_code: str = ""
    menu_path: str = ""
    
    # Design
    report_description: str = ""
    desain_report_description: str = ""
    selection_parameters: List[SelectionParameter] = field(default_factory=list)
    field_mappings: List[FieldMapping] = field(default_factory=list)
    validation_rules: List[str] = field(default_factory=list)
    special_processing: Dict[str, str] = field(default_factory=dict)

    # Additional dataset and lookup information
    valid_dataset_rules: List["DataConditionRow"] = field(default_factory=list)
    country_info: List["DataConditionRow"] = field(default_factory=list)
    currency_t500c: List["DataConditionRow"] = field(default_factory=list)
    currency_t001: List["DataConditionRow"] = field(default_factory=list)
    
    # Authorization
    authorization_objects: List[str] = field(default_factory=list)
    user_roles: List[str] = field(default_factory=list)
    
    # Design Constraints
    constraints: List[str] = field(default_factory=list)
    dependencies: List[str] = field(default_factory=list)
    
    # Error Handling
    error_scenarios: List[ErrorScenario] = field(default_factory=list)
    
    # Testing Requirements
    test_scenarios: List[TestScenario] = field(default_factory=list)
    test_data_location: str = ""
    
    # Design Changes
    design_changes: List[str] = field(default_factory=list)

# Fix 4: Improved dataclass_to_dict function with better error handling
def dataclass_to_dict(obj: Any) -> Any:
    """Recursively convert dataclasses and enums to plain Python types"""
    try:
        if is_dataclass(obj):
            obj = asdict(obj)
        if isinstance(obj, dict):
            return {k: dataclass_to_dict(v) for k, v in obj.items()}
        if isinstance(obj, list):
            return [dataclass_to_dict(v) for v in obj]
        if isinstance(obj, Enum):
            return obj.value
        # Handle other non-serializable types
        if hasattr(obj, '__dict__'):
            return dataclass_to_dict(obj.__dict__)
        return obj
    except Exception as e:
        logger.warning(f"Could not serialize object {type(obj)}: {e}")
        return str(obj)  # Fallback to string representation

# ================================
# CONFIGURATION MANAGEMENT
# ================================

class ConfigManager:
    """Manages configuration from environment variables and files"""
    
    def __init__(self, config_file: str = None):
        self.config = {}
        self._load_environment_config()
        if config_file and os.path.exists(config_file):
            self._load_file_config(config_file)
    
    def _load_environment_config(self):
        """Load configuration from environment variables"""
        self.config.update({
            'GEMINI_API_KEY': os.getenv('GEMINI_API_KEY'),
            'gemini_api_url': os.getenv('GEMINI_API_URL', 'https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-pro-latest:generateContent'),
            'project_name': os.getenv('PROJECT_NAME', 'System Integrator for Management Information System (MIS) Towards Single Source of Truth (SSoT)'),
            'default_output_dir': os.getenv('OUTPUT_DIR', './output'),
            'template_dir': os.getenv('TEMPLATE_DIR', './Template/FSD'),
            'default_output_dir': os.getenv('OUTPUT_DIR', '/Users/wahyu.perwira/Documents/Project/poc/SAP-AUTOMATE-FD-TD/backend/output/output'),
            'temperature': float(os.getenv('TEMPERATURE', '0.1')),
            'requirement_list_excel': os.getenv('REQUIREMENT_LIST_EXCEL', 'lookup-sheets/Requirement-List.xlsx')
        })
    
    def _load_file_config(self, config_file: str):
        """Load configuration from JSON file"""
        try:
            with open(config_file, 'r', encoding='utf-8') as f:
                file_config = json.load(f)
                self.config.update(file_config)
        except Exception as e:
            logger.warning(f"Could not load config file {config_file}: {e}")
    
    def get(self, key: str, default=None):
        """Get configuration value"""
        return self.config.get(key, default)
    
    def validate_required(self):
        """Validate required configuration"""
        if not os.getenv('GEMINI_API_KEY'):
            raise ValueError(
                "GEMINI_API_KEY is required. Set it as environment variable or in config file."
            )

# ================================
# MARKDOWN PARSER
# ================================

class MarkdownParser:
    """Parses markdown content and extracts structured data"""
    
    def __init__(self, markdown_content: str):
        self.markdown_content = markdown_content
        self.parsed_data = {}
        
    def parse(self) -> Dict[str, Any]:
        """Parse markdown content into structured data"""
        lines = self.markdown_content.split('\n')
        current_section = None
        current_subsection = None
        current_content = []
        
        for line in lines:
            line = line.strip()
            
            if line.startswith('# '):
                # Main title
                self.parsed_data['title'] = line[2:].strip()
            elif line.startswith('## '):
                # Save previous section
                if current_section:
                    self._save_section(current_section, current_subsection, current_content)
                
                # Start new section
                current_section = line[3:].strip()
                current_subsection = None
                current_content = []
            elif line.startswith('### '):
                # Save previous subsection
                if current_subsection:
                    self._save_subsection(current_section, current_subsection, current_content)
                
                # Start new subsection
                current_subsection = line[4:].strip()
                current_content = []
            elif line.startswith('|') and '|' in line:
                # Table row
                self._parse_table_row(line, current_section, current_subsection, current_content)
            elif line:
                # Regular content
                current_content.append(line)
        
        # Save final section
        if current_section:
            self._save_section(current_section, current_subsection, current_content)
        
        return self.parsed_data
    
    def _save_section(self, section: str, subsection: str, content: List[str]):
        """Save section content"""
        if section not in self.parsed_data:
            self.parsed_data[section] = {}
        
        if subsection:
            self._save_subsection(section, subsection, content)
        else:
            self.parsed_data[section]['content'] = content
    
    def _save_subsection(self, section: str, subsection: str, content: List[str]):
        """Save subsection content"""
        if section not in self.parsed_data:
            self.parsed_data[section] = {}
        
        if 'subsections' not in self.parsed_data[section]:
            self.parsed_data[section]['subsections'] = {}
        
        self.parsed_data[section]['subsections'][subsection] = content
    
    def _parse_table_row(self, line: str, section: str, subsection: str, content: List[str]):
        """Parse table row"""
        # Add table rows to content for now
        content.append(line)

# ================================
# WORD TEMPLATE PROCESSOR
# ================================

class WordTemplateProcessor:
    """Processes Word template documents and fills them with markdown data"""
    
    def __init__(self, template_path: str):
        self.template_path = template_path
        self.document = None
        self.template_doc = None
        
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template file not found: {template_path}")
    
    def load_template(self):
        """Load the Word template"""
        try:
            # Try loading as DocxTemplate first
            self.template_doc = DocxTemplate(self.template_path)
            logger.info(f"Loaded template as DocxTemplate: {self.template_path}")
        except Exception as e:
            logger.warning(f"Could not load as DocxTemplate: {e}")
            try:
                # Fallback to regular Document
                self.document = DocxDocument(self.template_path)
                logger.info(f"Loaded template as Document: {self.template_path}")
            except Exception as e2:
                logger.error(f"Could not load template: {e2}")
                raise
    
    def fill_template_with_markdown(self, markdown_content: str, output_path: str):
        """Fill template with markdown content"""
        # Parse markdown
        parser = MarkdownParser(markdown_content)
        parsed_data = parser.parse()
        
        if self.template_doc:
            self._fill_docx_template(parsed_data, output_path)
        else:
            self._fill_document_template(parsed_data, output_path)
    
    def _fill_docx_template(self, parsed_data: Dict[str, Any], output_path: str):
        """Fill DocxTemplate with parsed data"""
        try:
            # Prepare context data for template
            context = self._prepare_template_context(parsed_data)
            
            # Render template
            self.template_doc.render(context)
            
            # Save document
            self.template_doc.save(output_path)
            logger.info(f"Successfully generated document: {output_path}")
            
        except Exception as e:
            logger.error(f"Error filling DocxTemplate: {e}")
            # Fallback to manual filling
            self._manual_template_fill(parsed_data, output_path)
    
    def _fill_document_template(self, parsed_data: Dict[str, Any], output_path: str):
        """Fill regular Document with parsed data"""
        self._manual_template_fill(parsed_data, output_path)
    
    def _prepare_template_context(self, parsed_data: Dict[str, Any]) -> Dict[str, Any]:
        """Prepare context data for template rendering"""
        context = {
            'current_date': datetime.now().strftime('%d %B %Y'),
            'title': parsed_data.get('title', 'Functional Specification Design'),
            'program_name': '',
            'description': '',
            'project_name': 'System Integrator for Management Information System (MIS) Towards Single Source of Truth (SSoT)',
            'document_info': {},
            'requirements': {},
            'sap_objects': {},
            'design': {},
            'error_handling': {},
            'testing': {}
        }
        
        # Extract program name and description from title
        title_parts = parsed_data.get('title', '').split('\n')
        if len(title_parts) >= 2:
            context['program_name'] = title_parts[1].replace('## ', '').strip()
        if len(title_parts) >= 3:
            context['description'] = title_parts[2].replace('### ', '').strip()
        
        # Process sections
        for section_name, section_data in parsed_data.items():
            if section_name == 'title':
                continue
            
            section_key = self._normalize_section_name(section_name)
            
            if isinstance(section_data, dict):
                if 'content' in section_data:
                    context[section_key] = self._process_content(section_data['content'])
                
                if 'subsections' in section_data:
                    context[section_key] = {}
                    for subsection_name, subsection_content in section_data['subsections'].items():
                        subsection_key = self._normalize_section_name(subsection_name)
                        context[section_key][subsection_key] = self._process_content(subsection_content)
        
        return context
    
    def _normalize_section_name(self, name: str) -> str:
        """Normalize section name for template variables"""
        # Remove numbers and special characters, convert to lowercase
        normalized = re.sub(r'^\d+\.?\s*', '', name)  # Remove leading numbers
        normalized = re.sub(r'[^\w\s]', '', normalized)  # Remove special chars
        normalized = re.sub(r'\s+', '_', normalized.strip().lower())  # Replace spaces with underscores
        return normalized
    
    def _process_content(self, content: List[str]) -> Dict[str, Any]:
        """Process content list into structured data"""
        processed = {
            'text': [],
            'tables': [],
            'lists': []
        }
        
        current_table = []
        in_table = False
        
        for line in content:
            if line.startswith('|') and '|' in line:
                # Table row
                if not in_table:
                    in_table = True
                    current_table = []
                
                # Parse table row
                cells = [cell.strip() for cell in line.split('|')[1:-1]]  # Remove empty first/last
                current_table.append(cells)
            else:
                # End of table
                if in_table:
                    processed['tables'].append(current_table)
                    current_table = []
                    in_table = False
                
                if line.startswith('- '):
                    # List item
                    processed['lists'].append(line[2:].strip())
                elif line.startswith('**') and line.endswith('**'):
                    # Bold text
                    processed['text'].append({'type': 'bold', 'content': line[2:-2]})
                else:
                    # Regular text
                    if line:
                        processed['text'].append({'type': 'normal', 'content': line})
        
        # Don't forget final table
        if in_table and current_table:
            processed['tables'].append(current_table)
        
        return processed
    
    def _manual_template_fill(self, parsed_data: Dict[str, Any], output_path: str):
        """Manually fill template by replacing placeholders"""
        try:
            # Create new document based on template
            if self.document:
                doc = self.document
            else:
                doc = DocxDocument(self.template_path)
            
            # Replace placeholders in paragraphs
            self._replace_placeholders_in_document(doc, parsed_data)
            
            # Add content sections
            self._add_content_sections(doc, parsed_data)
            
            # Save document
            doc.save(output_path)
            logger.info(f"Successfully generated document using manual fill: {output_path}")
            
        except Exception as e:
            logger.error(f"Error in manual template fill: {e}")
            raise
    
    def _replace_placeholders_in_document(self, doc: DocxDocument, parsed_data: Dict[str, Any]):
        """Replace placeholders in document paragraphs"""
        # Common placeholders to replace
        placeholders = {
            '{{program_name}}': self._extract_program_name(parsed_data),
            '{{description}}': self._extract_description(parsed_data),
            '{{current_date}}': datetime.now().strftime('%d %B %Y'),
            '{{project_name}}': 'System Integrator for Management Information System (MIS) Towards Single Source of Truth (SSoT)'
        }
        
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for placeholder, value in placeholders.items():
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, value)
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for placeholder, value in placeholders.items():
                            if placeholder in paragraph.text:
                                paragraph.text = paragraph.text.replace(placeholder, value)
    
    def _add_content_sections(self, doc: DocxDocument, parsed_data: Dict[str, Any]):
        """Add content sections to document"""
        # Find insertion point (usually after a specific heading)
        insertion_point = len(doc.paragraphs)
        
        # Add each section
        for section_name, section_data in parsed_data.items():
            if section_name == 'title':
                continue
            
            # Add section heading
            heading = doc.add_heading(section_name, level=1)
            
            # Add section content
            if isinstance(section_data, dict):
                self._add_section_content(doc, section_data)
    
    def _add_section_content(self, doc: DocxDocument, section_data: Dict[str, Any]):
        """Add section content to document"""
        if 'content' in section_data:
            for content_line in section_data['content']:
                if content_line.strip():
                    doc.add_paragraph(content_line)
        
        if 'subsections' in section_data:
            for subsection_name, subsection_content in section_data['subsections'].items():
                # Add subsection heading
                doc.add_heading(subsection_name, level=2)
                
                # Check if content contains tables
                table_data = self._extract_table_from_content(subsection_content)
                if table_data:
                    self._add_table_to_document(doc, table_data)
                else:
                    # Add regular content
                    for content_line in subsection_content:
                        if content_line.strip() and not content_line.startswith('|'):
                            doc.add_paragraph(content_line)
    
    def _extract_table_from_content(self, content: List[str]) -> List[List[str]]:
        """Extract table data from content"""
        table_rows = []
        for line in content:
            if line.startswith('|') and line.count('|') >= 2:
                # Parse table row
                cells = [cell.strip() for cell in line.split('|')[1:-1]]
                if cells and not all(cell == '---' or cell.startswith('---') for cell in cells):
                    table_rows.append(cells)
        return table_rows
    
    def _add_table_to_document(self, doc: DocxDocument, table_data: List[List[str]]):
        """Add table to document"""
        if not table_data:
            return
        
        # Create table
        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
        table.style = 'Table Grid'
        
        # Fill table
        for row_idx, row_data in enumerate(table_data):
            for col_idx, cell_data in enumerate(row_data):
                if col_idx < len(table.rows[row_idx].cells):
                    table.rows[row_idx].cells[col_idx].text = cell_data
        
        # Style header row
        if table_data:
            for cell in table.rows[0].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
    
    def _extract_program_name(self, parsed_data: Dict[str, Any]) -> str:
        """Extract program name from parsed data"""
        title = parsed_data.get('title', '')
        lines = title.split('\n')
        for line in lines:
            if line.startswith('## '):
                return line[3:].strip()
        return 'Unknown Program'
    
    def _extract_description(self, parsed_data: Dict[str, Any]) -> str:
        """Extract description from parsed data"""
        title = parsed_data.get('title', '')
        lines = title.split('\n')
        for line in lines:
            if line.startswith('### '):
                return line[4:].strip()
        return 'Program Description'

# ================================
# FIXED MARKDOWN GENERATOR - ENSURES COMPLETE OUTPUT
# ================================

class EnhancedOutputGenerator:
    """Fixed output generator that ensures complete sections are always generated"""
    
    def __init__(self, config: ConfigManager):
        self.config = config
        self.output_dir = Path(config.get('default_output_dir'))
        self.template_dir = Path(config.get('template_dir'))
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def generate_all_outputs(self, fsd_document: FSDDocument, base_filename: str, 
                           template_path: str = None) -> Dict[str, str]:
        """Generate all output formats including Word document"""
        outputs = {}
        
        # Generate markdown first (single source of truth)
        md_file = self.output_dir / f"{base_filename}_fsd.md"
        markdown_content = self._generate_markdown(fsd_document)
        
        with open(md_file, 'w', encoding='utf-8') as f:
            f.write(markdown_content)
        outputs['markdown'] = str(md_file)
        
        # Generate JSON
        json_file = self.output_dir / f"{base_filename}_fsd.json"
        outputs['json'] = str(json_file)
        self._generate_json(fsd_document, json_file)
        
        # Generate Word document from markdown
        if template_path and os.path.exists(template_path):
            docx_file = self.output_dir / f"{base_filename}_fsd.docx"
            outputs['docx'] = str(docx_file)
            self._generate_word_document(markdown_content, template_path, docx_file)
        else:
            # Try to find template in template directory
            template_file = self._find_template_file()
            if template_file:
                docx_file = self.output_dir / f"{base_filename}_fsd.docx"
                outputs['docx'] = str(docx_file)
                self._generate_word_document(markdown_content, template_file, docx_file)
            else:
                logger.warning("No template file found, skipping Word document generation")
        
        # Generate summary
        summary_file = self.output_dir / f"{base_filename}_summary.txt"
        outputs['summary'] = str(summary_file)
        self._generate_summary(fsd_document, summary_file)
        
        return outputs
    
    def _find_template_file(self) -> Optional[str]:
        """Find template file in template directory"""
        if not self.template_dir.exists():
            return None
        
        # Look for .docx files in template directory
        for file_path in self.template_dir.glob("*.docx"):
            if 'FSD' in file_path.name or 'template' in file_path.name.lower():
                return str(file_path)
        
        # If no specific FSD template found, use any .docx file
        docx_files = list(self.template_dir.glob("*.docx"))
        if docx_files:
            return str(docx_files[0])
        
        return None
    
    def _generate_word_document(self, markdown_content: str, template_path: str, output_file: Path):
        """Generate Word document from markdown using template"""
        try:
            processor = WordTemplateProcessor(template_path)
            processor.load_template()
            processor.fill_template_with_markdown(markdown_content, str(output_file))
            logger.info(f"Generated Word document: {output_file}")
        except Exception as e:
            logger.error(f"Failed to generate Word document: {e}")
    
    def _lookup_assign_nodin(self, ricefw_id: str) -> Optional[str]:
        """Lookup Assign Nodin value for a given RICEFW ID from Excel sheet"""
        excel_path = Path(self.config.get('requirement_list_excel'))
        if not excel_path.exists():
            return None
        try:
            wb = openpyxl.load_workbook(excel_path, data_only=True)
            ws = wb.active
            headers = [cell.value for cell in next(ws.iter_rows(min_row=1, max_row=1))]
            id_idx = headers.index('SAP WRICEF ID')
            nodin_idx = headers.index('Assign Nodin')
            for row in ws.iter_rows(min_row=2, values_only=True):
                if str(row[id_idx]).strip() == ricefw_id:
                    return str(row[nodin_idx]).strip() if row[nodin_idx] is not None else None
        except Exception as exc:
            logger.warning(f"Failed to lookup Assign Nodin for {ricefw_id}: {exc}")
        return None
    
    def _lookup_requirement_description(self, ricefw_id: str) -> Optional[str]:
        """Lookup Requirement Description for a given RICEFW ID from Excel sheet"""
        excel_path = Path(self.config.get('requirement_list_excel'))
        if not excel_path.exists():
            return None
        try:
            wb = openpyxl.load_workbook(excel_path, data_only=True)
            ws = wb.active
            headers = [cell.value for cell in next(ws.iter_rows(min_row=1, max_row=1))]
            id_idx = headers.index('SAP WRICEF ID')
            desc_idx = headers.index('Requirement Description')
            for row in ws.iter_rows(min_row=2, values_only=True):
                if str(row[id_idx]).strip() == ricefw_id:
                    return str(row[desc_idx]).strip() if row[desc_idx] is not None else None
        except Exception as exc:
            logger.warning(f"Failed to lookup Requirement Description for {ricefw_id}: {exc}")
        return None

    def _improve_text(self, text: str) -> str:
        """Use LLM to improve Indonesian phrasing"""
        async def _run(text_val: str) -> str:
            async with LLMClient(self.config) as llm:
                prompt = (
                    "Sempurnakan kalimat berikut agar lebih jelas dan profesional tetang kebutuhan pengguna SAP dalam Bahasa Indonesia dengan menjabarkan detail yang lengkap. "
                    "Kembalikan JSON {\"result\": \"kalimat\"}.\n\nKalimat: " + text_val
                )
                result = await llm.analyze(prompt, {"task": "improve_requirement"})
                return result.get('result', text_val)

        try:
            return asyncio.run(_run(text))
        except Exception as exc:
            logger.warning(f"Failed to improve text via LLM: {exc}")
            return text
        
    def _compose_report_description(self, fsd_document: FSDDocument) -> str:
        """Generate a fallback report description if none provided"""
        if fsd_document.desain_report_description:
            return fsd_document.desain_report_description

        base_desc = ""
        # Try to use requirement description from Excel
        for doc in fsd_document.related_documents:
            if doc.startswith("RICEFW ID:"):
                ricefw_id = doc.split(":", 1)[1].strip()
                lookup_desc = self._lookup_requirement_description(ricefw_id)
                if lookup_desc:
                    base_desc = lookup_desc
                break

        if not base_desc:
            base_desc = fsd_document.user_requirements

        parts = [base_desc] if base_desc else []
        if fsd_document.selection_parameters:
            parts.append(
                f"Report ini memiliki {len(fsd_document.selection_parameters)} parameter selection screen."
            )
        if fsd_document.field_mappings:
            parts.append(
                f"Proses detail mencakup {len(fsd_document.field_mappings)} field utama."
            )

        if not parts:
            parts.append(
                "Report ini dirancang untuk memenuhi kebutuhan bisnis yang telah ditentukan."
            )

        description = " ".join(parts)
        return self._improve_text(description)
            
    def _generate_markdown(self, fsd_document: FSDDocument) -> str:
        """Generate Markdown content"""
        md_lines = []
        
        # Title
        md_lines.extend([
            f"# Functional Specification Design (FSD)",
            f"## {fsd_document.program_name}",
            f"### {fsd_document.report_description}",
            "",
            f"**Generated on:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            ""
        ])
        
        # Document Information
        md_lines.extend([
            "## 1. INFORMASI DOKUMEN",
            "",
            f"- **Project**: {fsd_document.project_name}",
            f"- **Document Location**: {fsd_document.document_location}",
            ""
        ])
        
        if fsd_document.related_documents:
            md_lines.append("- **Related Documents**:")
            for doc in fsd_document.related_documents:
                # md_lines.append(f"  - {doc}") - innitial code
                if doc.startswith("RICEFW ID:"):
                    ricefw_id = doc.split(":", 1)[1].strip()
                    nodin = self._lookup_assign_nodin(ricefw_id)
                    if nodin:
                        md_lines.append(f"  - RICEFW ID: {nodin}")
                    else:
                        md_lines.append(f"  - {doc}")
                else:
                    md_lines.append(f"  - {doc}")
            md_lines.append("")
        
        if fsd_document.reviewers:
            md_lines.append("- **Reviewers**:")
            for reviewer in fsd_document.reviewers:
                md_lines.append(f"  - {reviewer.get('role', '')}: {reviewer.get('name', '')}")
            md_lines.append("")
        
        if fsd_document.version_history:
            md_lines.append("- **Version History**:")
            for version in fsd_document.version_history:
                md_lines.append(f"  - {version.get('version', '')}: {version.get('change', '')} by {version.get('author', '')} on {version.get('date', '')}")
            md_lines.append("")
        
        # Determine user requirement from Excel if possible
        requirement_value = fsd_document.user_requirements
        for doc in fsd_document.related_documents:
            if doc.startswith("RICEFW ID:"):
                ricefw_id = doc.split(":", 1)[1].strip()
                desc = self._lookup_requirement_description(ricefw_id)
                if desc:
                    requirement_value = self._improve_text(desc)
                break
        
        # General Requirements
        md_lines.extend([
            "## 2. PERSYARATAN UMUM",
            "",
            # f"**User Requirements**: {fsd_document.user_requirements}",
            f"**User Requirements**: {requirement_value}",
            "",
            "**Assumptions**:",
        ])
        
        for assumption in fsd_document.assumptions:
            md_lines.append(f"- {assumption}")
        md_lines.append("")

        # Existing SAP Objects - for loop from the excell information
        program_name_value = fsd_document.program_name
        for doc in fsd_document.related_documents:
            if doc.startswith("RICEFW ID:"):
                ricefw_id = doc.split(":", 1)[1].strip()
                nodin_val = self._lookup_assign_nodin(ricefw_id)
                if nodin_val:
                    program_name_value = nodin_val
                break

        
        # Existing SAP Objects
        md_lines.extend([
            "## 3. EXISTING SAP OBJECTS",
            "",
            # f"- **Program Name**: {fsd_document.program_name}",
            f"- **Program Name**: {program_name_value}",
            f"- **Transaction Code**: {fsd_document.transaction_code}",
            f"- **Menu Path**: {fsd_document.menu_path}",
            ""
        ])

        #Additional Design Section
        desain_description_value = self._compose_report_description(fsd_document)
        md_lines.extend([
            "## 4. DESAIN",
            "",
            "### 4.1 Description detail dari Report",
            "",
            desain_description_value,
            ""
        ])
        
        # Design - Selection Screen
        if fsd_document.selection_parameters:
            md_lines.extend([
                "### 4.2 Selection Screen",
                "",
                "| Parameter | Type | Description | Mandatory | Select-Option | No Intervals |",
                "|-----------|------|-------------|-----------|---------------|--------------|"
            ])
            
            for param in fsd_document.selection_parameters:
                mandatory = "Yes" if param.is_mandatory else "No"
                select_opt = "Yes" if param.is_select_option else "No"
                no_intervals = "Yes" if param.has_no_intervals else "No"
                md_lines.append(f"| {param.name} | {param.type} | {param.description} | {mandatory} | {select_opt} | {no_intervals} |")
            
            md_lines.append("")
        
        # Design - Detail Processing
        if fsd_document.field_mappings:
            md_lines.extend([
                "### 4.3 Detail Processing",
                "",
                "| Field Name | Technical Field | Source Table | Processing Logic | Processing Type |",
                "|------------|-----------------|--------------|------------------|-----------------|"
            ])
            
            for field in fsd_document.field_mappings:
                processing_type = field.processing_type.value if isinstance(field.processing_type, Enum) else str(field.processing_type)
                md_lines.append(f"| {field.display_name} | {field.technical_field} | {field.source_table} | {field.processing_logic} | {processing_type} |")
            
            md_lines.append("")

        if fsd_document.valid_dataset_rules:
            md_lines.extend([
                "### 4.4 Detail Process Only valid datasets",
                "",
                "| Data | Kondisi |",
                "|------|---------|",
            ])

            for row in fsd_document.valid_dataset_rules:
                md_lines.append(f"| {row.data} | {row.condition} |")

            md_lines.append("")

        def append_form(title: str, rows: List[DataConditionRow]):
            if not rows:
                return
            md_lines.extend([
                f"### {title}",
                "",
                "| Data | Kondisi |",
                "|------|---------|",
            ])
            for r in rows:
                md_lines.append(f"| {r.data} | {r.condition} |")
            md_lines.append("")

        append_form("Form Get_Country_Info", fsd_document.country_info)
        append_form("Form Get_Currency_T500C", fsd_document.currency_t500c)
        append_form("Form Get_Currency_T001", fsd_document.currency_t001)
    
            # Error Handling
        if fsd_document.error_scenarios:
            md_lines.extend([
                "## 5. PENANGANAN ERROR",
                "",
                "| No | Error Description | Resolution | Error Code | Severity |",
                "|----|-------------------|------------|------------|----------|"
            ])
            
            for i, error in enumerate(fsd_document.error_scenarios, 1):
                md_lines.append(f"| {i} | {error.error_description} | {error.resolution} | {error.error_code} | {error.severity} |")
            
            md_lines.append("")
        
        # Testing Requirements
        if fsd_document.test_scenarios:
            md_lines.extend([
                "## 6. PERSYARATAN PENGUJIAN",
                "",
                "| No | Test Condition | Expected Result | Test Data | Priority |",
                "|----|----------------|-----------------|-----------|----------|"
            ])
            
            for i, test in enumerate(fsd_document.test_scenarios, 1):
                md_lines.append(f"| {i} | {test.condition} | {test.expected_result} | {test.test_data} | {test.priority} |")
        
        return '\n'.join(md_lines)
    

    def _generate_json(self, fsd_document: FSDDocument, output_file: Path):
        """Generate JSON output"""
        # Use the existing dataclass_to_dict function instead of custom conversion
        doc_dict = dataclass_to_dict(fsd_document)
        
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(doc_dict, f, indent=2, ensure_ascii=False)
        
        logger.info(f"Generated JSON output: {output_file}")

    
    def _generate_summary(self, fsd_document: FSDDocument, output_file: Path):
        """Generate summary report"""
        summary_lines = [
            "=== FSD ANALYSIS SUMMARY ===",
            f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            "",
            f"Program Name: {fsd_document.program_name}",
            f"Description: {fsd_document.report_description}",
            # f"Description: {self._compose_report_description(fsd_document)}",
            f"Transaction Code: {fsd_document.transaction_code}",
            "",
            "STATISTICS:",
            f"- Selection Parameters: {len(fsd_document.selection_parameters)}",
            f"- Field Mappings: {len(fsd_document.field_mappings)}",
            f"- Error Scenarios: {len(fsd_document.error_scenarios)}",
            f"- Test Scenarios: {len(fsd_document.test_scenarios)}",
            f"- Validation Rules: {len(fsd_document.validation_rules)}",
            f"- Authorization Objects: {len(fsd_document.authorization_objects)}",
            "",
            "ANALYSIS COMPLETENESS:",
            f"- User Requirements: {'✓' if fsd_document.user_requirements else '✗'}",
            f"- Assumptions: {'✓' if fsd_document.assumptions else '✗'}",
            f"- Field Mappings: {'✓' if fsd_document.field_mappings else '✗'}",
            f"- Error Handling: {'✓' if fsd_document.error_scenarios else '✗'}",
            f"- Test Coverage: {'✓' if fsd_document.test_scenarios else '✗'}"
        ]
        
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write('\n'.join(summary_lines))
        
        logger.info(f"Generated summary report: {output_file}")

# ================================
# LLM CLIENT
# ================================

class LLMClient:
    """Asynchronous LLM client for Gemini API"""
    
    def __init__(self, config: ConfigManager):
        self.config = config
        self.session = None
        # self.api_key = os.getenv('GEMINI_API_KEY')
        # self.api_url = config.get('gemini_api_url')
        self.api_key = os.getenv('GEMINI_API_KEY', '').strip('"')
        self.api_url = config.get('gemini_api_url', '').strip('"')
        self.max_tokens = config.get('max_tokens')
        self.temperature = config.get('temperature')
    
    async def __aenter__(self):
        self.session = aiohttp.ClientSession()
        return self
    
    async def __aexit__(self, exc_type, exc_val, exc_tb):
        if self.session:
            await self.session.close()
    
    async def analyze(self, prompt: str, context: Dict[str, Any] = None) -> Dict[str, Any]:
        """Send analysis request to LLM"""
        try:
            full_prompt = self._build_prompt(prompt, context)
            response = await self._call_api(full_prompt)
            return self._parse_response(response)
        except Exception as e:
            logger.error(f"LLM analysis failed: {e}")
            return {}
    
    def _build_prompt(self, prompt: str, context: Dict[str, Any] = None) -> str:
        system_prompt = """
        Anda adalah seorang pengembang SAP ABAP sekaligus konsultan fungsional yang sangat berpengalaman dalam membuat
        dokumen Functional Specification Design (FSD). Anda menganalisis kode ABAP dan mengekstrak informasi penting untuk
        menyusun dokumentasi FSD yang komprehensif.

        Tugas Anda adalah:
        1. Menganalisis kode ABAP beserta komentar yang diberikan.
        2. Menghasilkan informasi terstruktur untuk setiap bagian FSD.

        Ketentuan:
        - Selalu berikan respons dalam format JSON yang valid.
        - Analisis harus teliti dan akurat.
        - Anda **WAJIB** menjawab **HANYA** dalam **Bahasa Indonesia**; dilarang menggunakan bahasa lain.
        """
        
        context_section = ""
        if context:
            context_section = f"\n\nContext Information:\n{json.dumps(context, indent=2)}"
        
        return f"{system_prompt}\n\n{prompt}{context_section}"
    
    async def _call_api(self, prompt: str) -> Dict[str, Any]:
        """Make API call to Gemini"""
        headers = {
            "Content-Type": "application/json"
        }
        
        payload = {
            "contents": [
                {
                    "parts": [
                        {"text": prompt}
                    ]
                }
            ],
            "generationConfig": {
                "temperature": self.temperature,
                "topK": 1,
                "topP": 1,
                "maxOutputTokens": self.max_tokens,
                "response_mime_type": "application/json"
            }
        }
        
        url = f"{self.api_url}?key={self.api_key}"
        print(f"[Debug - Call API] Calling LLM API: {url}")
        # print(f"[Debug - Call API] Payload: {json.dumps(payload, indent=2)}")
        
        async with self.session.post(url, headers=headers, json=payload) as response:
            if response.status == 200:
                result = await response.json()
                return result
            else:
                error_text = await response.text()
                raise Exception(f"API call failed with status {response.status}: {error_text}")
    
    def _parse_response(self, response: Dict[str, Any]) -> Dict[str, Any]:
        """Parse API response and extract JSON content"""
        try:
            content = response.get("candidates", [{}])[0].get("content", {}).get("parts", [{}])[0].get("text", "{}")
            return json.loads(content)
        except (json.JSONDecodeError, KeyError) as e:
            logger.error(f"Failed to parse LLM response: {e}")
            return {}

# ================================
# ABAP HTML EXTRACTOR
# ================================

class ABAPHTMLExtractor:
    """Enhanced ABAP HTML extractor"""
    
    def __init__(self, html_file_path: str):
        self.file_path = html_file_path
        self.soup = None
        self.raw_data = {}
    
    def extract_all(self) -> Dict[str, Any]:
        """Extract all information from HTML file"""
        if not os.path.exists(self.file_path):
            raise FileNotFoundError(f"File not found: {self.file_path}")
        
        with open(self.file_path, 'r', encoding='utf-8') as file:
            html_content = file.read()
        
        self.soup = BeautifulSoup(html_content, 'html.parser')
        
        # Extract raw content
        code_divs = self.soup.find_all('div', class_='code')
        comment_divs = self.soup.find_all('div', class_='codeComment')
        
        self.raw_data = {
            'file_path': self.file_path,
            'file_name': os.path.basename(self.file_path),
            'raw_code': '\n'.join([div.get_text() for div in code_divs]),
            'raw_comments': '\n'.join([div.get_text() for div in comment_divs]),
            'html_title': self._extract_title(),
            'html_description': self._extract_description()
        }
        
        return self.raw_data
    
    def _extract_title(self) -> str:
        """Extract title from HTML"""
        title_elem = self.soup.find('title')
        if title_elem:
            return title_elem.get_text().strip()
        
        h2_elem = self.soup.find('h2')
        if h2_elem:
            return h2_elem.get_text().replace('Code listing for: ', '').strip()
        
        return ""
    
    def _extract_description(self) -> str:
        """Extract description from HTML"""
        h3_elem = self.soup.find('h3')
        if h3_elem:
            return h3_elem.get_text().replace('Description: ', '').strip()
        return ""

class IntelligentFSDMapper:
    """Fixed Intelligent FSD mapper with comprehensive structure analysis"""
    
    def __init__(self, config: ConfigManager):
        self.config = config
        self.fsd_document = FSDDocument()
        self.fsd_document.project_name = config.get('project_name')
        
    async def analyze_and_map(self, html_file_path: str) -> FSDDocument:
        """Main method to analyze HTML and create FSD document"""
        logger.info(f"Starting fixed comprehensive FSD mapping for: {html_file_path}")
        
        # Extract raw data from HTML
        extractor = ABAPHTMLExtractor(html_file_path)
        raw_data = extractor.extract_all()
        
        # Set basic document info
        self.fsd_document.document_location = raw_data['file_name']
        
        # Fixed comprehensive LLM analysis
        async with LLMClient(self.config) as llm:
            await self._analyze_with_fixed_comprehensive_llm(llm, raw_data)
        
        logger.info("Fixed comprehensive FSD mapping completed successfully")
        return self.fsd_document
    
    async def _analyze_with_fixed_comprehensive_llm(self, llm: LLMClient, raw_data: Dict[str, Any]):
        """Perform fixed comprehensive LLM analysis"""
        
        # Enhanced analysis tasks with FIXED prompts
        analysis_tasks = [
            ("basic_info", self._create_enhanced_basic_info_prompt(raw_data)),
            ("selection_screen", self._create_enhanced_selection_screen_prompt(raw_data)),
            ("complete_field_mappings", self._create_complete_field_mappings_prompt(raw_data)),  # FIXED
            ("complete_valid_datasets", self._create_complete_valid_datasets_prompt(raw_data)),  # FIXED
            ("complete_lookup_forms", self._create_complete_lookup_forms_prompt(raw_data)),      # FIXED
            ("error_handling", self._create_error_handling_prompt(raw_data)),
            ("test_scenarios", self._create_test_scenarios_prompt(raw_data)),
            ("validation_rules", self._create_validation_rules_prompt(raw_data)),
            ("authorization", self._create_authorization_prompt(raw_data))
        ]
        
        # Execute all analyses
        results = {}
        for task_name, prompt in analysis_tasks:
            try:
                logger.info(f"Analyzing with fixed prompt: {task_name}")
                result = await llm.analyze(prompt, {"task": task_name})
                results[task_name] = result
            except Exception as e:
                logger.error(f"Failed to analyze {task_name}: {e}")
                results[task_name] = {}
        
        # Map results to FSD document
        await self._map_fixed_results_to_fsd(results)
    
    def _create_complete_field_mappings_prompt(self, raw_data: Dict[str, Any]) -> str:
        """FIXED: Complete field mappings prompt that captures ALL structure fields"""
        return f"""
        Analisis kode ABAP secara MENYELURUH untuk mengekstrak SEMUA field dari struktur output report.

        TUGAS ANALISIS SANGAT KOMPREHENSIF:
        
        **LANGKAH 1: IDENTIFIKASI STRUKTUR DATA UTAMA**
        - Cari struktur utama yang digunakan untuk output (gty_report, lty_hr_report, dll)
        - Bisa berada di bagian global TYPES: atau dalam CLASS DEFINITION
        - Ekstrak SEMUA field dari struktur tersebut dengan tipe datanya
        
        **LANGKAH 2: ANALISIS ASSIGNMENT LOGIC**
        - Analisis method get_data, f_get_data, atau method utama untuk pengisian data
        - Perhatikan SETIAP assignment (lw_result-field = source, lw_report-field = source)
        - Ekstrak SEMUA operasi READ table untuk lookup
        
        **LANGKAH 3: ANALISIS ALV COLUMN DEFINITION**
        - Cari set_short_text, set_medium_text, set_long_text dalam method display_alv atau f_display_alv
        - Ekstrak column names dan display text yang didefinisikan
        
        **LANGKAH 4: IDENTIFIKASI AGGREGATION**
        - Cari add_aggregation calls untuk mengetahui field mana yang di-SUM atau COUNT
        
        **POLA STRUKTUR YANG HARUS DICARI:**
        ```
        TYPES: BEGIN OF gty_report,  
               atau
        TYPES: BEGIN OF lty_hr_report,
        ```
        
        **POLA ASSIGNMENT YANG HARUS DICARI:**
        ```
        lw_report-bukrs = p0001-bukrs.
        lw_result-butxt = <fs_butxt>-butxt.
        READ TABLE lt_t001 ... WITH KEY bukrs = ...
        ```
        
        **PASTIKAN MENGANALISIS SEMUA FIELD INI JIKA ADA:**
        - bukrs, butxt (company code & text)
        - gsber, gtext/gsber_text (business area & text)  
        - count (counter)
        - pernr (personnel number)
        - pnalt/prev_pernr (previous personnel number)
        - name/ename (employee name)
        - begda (start date)
        - lgart (wage type)
        - lgtxt (wage type text)
        - betrg (amount)
        - waers (currency)
        - ocrsn (reason)
        - ocrtx (reason text)
        
        Kembalikan JSON dengan format:
        {{
            "field_mappings": [
                {{
                    "display_name": "nama_tampilan_dari_ALV_atau_inferensi_yang_tepat",
                    "technical_field": "nama_field_teknis_dari_struktur",
                    "source_table": "tabel_sumber_data_yang_sebenarnya",
                    "processing_logic": "penjelasan_sangat_detail_bagaimana_field_diisi_dengan_assignment_atau_lookup_yang_spesifik",
                    "processing_type": "DIRECT|LOOKUP|CALCULATION|CONSTANT|AGGREGATION",
                    "join_condition": "kondisi_join_dari_read_table_jika_ada",
                    "where_condition": "kondisi_where_dari_select_jika_ada",
                    "is_aggregated": true/false,
                    "aggregation_type": "SUM|COUNT jika ada aggregation"
                }}
            ]
        }}

        **WAJIB: ANALISIS HARUS MENCAKUP SEMUA FIELD YANG ADA DI STRUKTUR DATA!**
        **JANGAN LEWATKAN FIELD APAPUN!**

        Kode ABAP:
        {raw_data.get('raw_code', '')}

        Komentar:
        {raw_data.get('raw_comments', '')}
        """
    
    def _create_complete_valid_datasets_prompt(self, raw_data: Dict[str, Any]) -> str:
        """FIXED: Complete valid datasets prompt that analyzes actual validation logic"""
        return f"""
        Analisis kode ABAP secara MENYELURUH untuk mengidentifikasi aturan dataset valid dan logika validasi.

        TUGAS ANALISIS SANGAT KOMPREHENSIF:
        
        **LANGKAH 1: IDENTIFIKASI INFOTYPES**
        - Cari deklarasi INFOTYPES: di awal program
        - Identifikasi semua infotype yang dideklarasikan (0001, 0015, 0032, 0267, dll)
        
        **LANGKAH 2: ANALISIS PEMBACAAN INFOTYPE**
        - Cari penggunaan rp_read_infotype untuk setiap infotype
        - Analisis periode yang digunakan (gv_begda, gv_endda, pn-begda, pn-endda)
        
        **LANGKAH 3: ANALISIS VALIDASI DATA**
        - Cari method validation, f_validation, process_data
        - Analisis semua DELETE ... WHERE ... statements
        - Identifikasi kondisi validasi untuk setiap dataset
        
        **LANGKAH 4: ANALISIS LOGIKA PERIODE**
        - Cari method populate_range_date, get_populate_range_date
        - Analisis bagaimana periode ditentukan
        
        **POLA YANG HARUS DICARI:**
        ```
        INFOTYPES: 0001 MODE N,
                   0015 MODE N,
                   0032 MODE N,
                   0267 MODE N.
        
        rp_read_infotype peras-pernr 0001 p0001 ...
        
        DELETE gt_report WHERE lgart NOT IN s_lgart.
        DELETE lt_hr_report WHERE ...
        ```
        
        **LOGIKA VALIDASI YANG HARUS DIANALISIS:**
        - Validasi wage type berdasarkan selection screen
        - Validasi date range dan periode
        - Validasi reason codes
        - Logika pembersihan data invalid
        
        Kembalikan JSON dengan format:
        {{
            "valid_dataset_rules": [
                {{
                    "data": "nama_infotype_atau_dataset_spesifik_seperti_P0001_P0015_dll",
                    "condition": "deskripsi_kondisi_validasi_yang_sangat_detail_berdasarkan_analisis_kode_actual_termasuk_DELETE_WHERE_logic",
                    "validation_method": "nama_method_yang_melakukan_validasi",
                    "period_logic": "logika_periode_yang_digunakan_untuk_membaca_data",
                    "infotype_usage": "bagaimana_infotype_ini_digunakan_dalam_program"
                }}
            ]
        }}

        **WAJIB: HARUS MENGANALISIS SEMUA INFOTYPE YANG DIDEKLARASIKAN!**
        **KONDISI HARUS BERDASARKAN KODE ACTUAL, BUKAN TEMPLATE GENERIC!**

        Kode ABAP:
        {raw_data.get('raw_code', '')}

        Komentar:
        {raw_data.get('raw_comments', '')}
        """
    
    def _create_complete_lookup_forms_prompt(self, raw_data: Dict[str, Any]) -> str:
        """FIXED: Complete lookup forms prompt that analyzes all lookup operations"""
        return f"""
        Analisis kode ABAP secara MENYELURUH untuk mengidentifikasi SEMUA operasi lookup dan helper logic.

        TUGAS ANALISIS SANGAT KOMPREHENSIF:
        
        **LANGKAH 1: IDENTIFIKASI MASTER DATA LOADING**
        - Cari method get_additional_data, f_get_dtl, atau method untuk load master data
        - Analisis semua SELECT statements ke tabel master (T001, TGSBT, T512T, T52OCRT, T500P, dll)
        
        **LANGKAH 2: ANALISIS LOOKUP OPERATIONS**
        - Cari semua READ table operations dalam method get_data atau f_get_data
        - Identifikasi key yang digunakan untuk lookup
        - Analisis sy-subrc checking setelah read table
        
        **LANGKAH 3: KATEGORISASI LOOKUP BERDASARKAN FUNGSI**
        - Company Code lookup (T001)
        - Business Area lookup (TGSBT)
        - Wage Type lookup (T512T)
        - Payroll Reason lookup (T52OCRT)
        - Country/Personnel Area lookup (T500P)
        - Currency lookup (berbagai tabel currency)
        
        **LANGKAH 4: ANALISIS ERROR HANDLING**
        - Bagaimana program menangani lookup yang gagal
        - Default values atau empty handling
        
        **POLA YANG HARUS DICARI:**
        ```
        SELECT bukrs butxt FROM t001 INTO TABLE lt_t001.
        
        READ TABLE lt_t001 ASSIGNING <fs_butxt> WITH KEY bukrs = p0001-bukrs BINARY SEARCH.
        IF sy-subrc EQ 0.
          lw_result-butxt = <fs_butxt>-butxt.
        ENDIF.
        ```
        
        **TABEL LOOKUP YANG HARUS DIANALISIS:**
        - T001: Company Code master
        - TGSBT: Business Area text
        - T512T: Wage Type text
        - T52OCRT: Payroll reason text
        - T500P: Personnel area (jika ada)
        - T500C: Currency by country (jika ada)
        
        Kembalikan JSON dengan format:
        {{
            "company_code_lookup": [
                {{
                    "source_table": "T001",
                    "target_fields": ["BUTXT"],
                    "lookup_key": "BUKRS",
                    "condition": "Ambil BUTXT dari tabel T001 dengan key BUKRS = PA0001-BUKRS. Jika sy-subrc = 0, isi field butxt dengan hasil lookup.",
                    "error_handling": "Jika lookup gagal, field butxt tetap kosong",
                    "method_location": "method_get_additional_data_dan_get_data"
                }}
            ],
            "business_area_lookup": [
                {{
                    "source_table": "TGSBT",
                    "target_fields": ["GTEXT"],
                    "lookup_key": "GSBER",
                    "condition": "Ambil GTEXT dari tabel TGSBT dengan key GSBER = PA0001-GSBER dan SPRAS = 'E'. Jika sy-subrc = 0, isi field gtext/gsber_text.",
                    "error_handling": "Jika lookup gagal, field gtext tetap kosong",
                    "method_location": "method_get_additional_data_dan_get_data"
                }}
            ],
            "wage_type_lookup": [
                {{
                    "source_table": "T512T",
                    "target_fields": ["LGTXT"],
                    "lookup_key": "LGART",
                    "condition": "Ambil LGTXT dari tabel T512T dengan key LGART = PA0015/PA0267-LGART, SPRSL = 'E', dan MOLGA = '34'. Jika sy-subrc = 0, isi field lgtxt.",
                    "error_handling": "Jika lookup gagal, field lgtxt tetap kosong",
                    "method_location": "method_get_additional_data_dan_get_data"
                }}
            ],
            "payroll_reason_lookup": [
                {{
                    "source_table": "T52OCRT",
                    "target_fields": ["OCRTX"],
                    "lookup_key": "OCRSN",
                    "condition": "Ambil OCRTX dari tabel T52OCRT dengan key OCRSN = PA0267-OCRSN, SPRSL = 'E', dan MOLGA = '34'. Jika sy-subrc = 0, isi field ocrtx.",
                    "error_handling": "Jika lookup gagal, field ocrtx tetap kosong", 
                    "method_location": "method_get_additional_data_dan_get_data"
                }}
            ],
            "other_lookups": [
                {{
                    "lookup_name": "nama_lookup_lainnya_jika_ada",
                    "source_table": "tabel_sumber",
                    "target_fields": ["field_target"],
                    "lookup_key": "key_field",
                    "condition": "deskripsi_lengkap_logic_lookup",
                    "error_handling": "penanganan_error",
                    "method_location": "lokasi_method"
                }}
            ]
        }}

        **WAJIB: HARUS MENGANALISIS SEMUA READ TABLE OPERATIONS!**
        **KONDISI HARUS MENJELASKAN LOGIC ACTUAL DARI KODE!**

        Kode ABAP:
        {raw_data.get('raw_code', '')}

        Komentar:
        {raw_data.get('raw_comments', '')}
        """
    
    # Keep other existing prompts but enhance them
    def _create_enhanced_basic_info_prompt(self, raw_data: Dict[str, Any]) -> str:
        """Enhanced basic info prompt"""
        return f"""
        Analisis kode ABAP secara mendalam untuk mengekstrak semua informasi dasar program.

        TUGAS ANALISIS KOMPREHENSIF:
        1. Identifikasi nama program dari REPORT statement
        2. Ekstrak deskripsi lengkap dari komentar header dan HTML description
        3. Cari RICEFW ID, tanggal pembuatan, pembuat program dari komentar
        4. Analisis tujuan bisnis program dari nama, deskripsi, dan logika kode
        5. Identifikasi asumsi dari komentar dan logika program
        6. Tentukan transaction code berdasarkan nama program (jika mengikuti pattern)
        7. Cari functional contact dan informasi tim dari komentar

        Kembalikan JSON dengan format:
        {{
            "program_name": "nama program dari REPORT statement",
            "report_description": "deskripsi singkat dan jelas dari HTML description atau komentar",
            "desain_report_description": "deskripsi teknis detail lengkap tentang arsitektur dan desain program, cara kerja program, tabel yang digunakan, flow pemrosesan data, dan output yang dihasilkan",
            "user_requirements": "kebutuhan pengguna yang disimpulkan dari tujuan dan fungsi program",
            "assumptions": ["daftar asumsi lengkap berdasarkan komentar dan logika program"],
            "transaction_code": "kode transaksi jika ada atau dapat diturunkan dari nama program",
            "menu_path": "alur menu jika ada dalam komentar, jika tidak ada tulis 'N/A'",
            "ricefw_id": "RICEFW ID dari komentar header",
            "created_date": "tanggal pembuatan dari komentar",
            "created_by": "nama pembuat program dari komentar",
            "functional_contact": "nama kontak fungsional dari komentar"
        }}

        Kode ABAP:
        {raw_data.get('raw_code', '')}

        Komentar:
        {raw_data.get('raw_comments', '')}

        HTML Title: {raw_data.get('html_title', '')}
        HTML Description: {raw_data.get('html_description', '')}
        """
    
    def _create_enhanced_selection_screen_prompt(self, raw_data: Dict[str, Any]) -> str:
        """Enhanced selection screen prompt"""
        return f"""
        Analisis kode ABAP secara komprehensif untuk semua elemen selection screen.

        TUGAS ANALISIS KOMPREHENSIF:
        1. Identifikasi semua PARAMETERS dan SELECT-OPTIONS
        2. Analisis SELECTION-SCREEN BLOCK untuk grouping
        3. Ekstrak tipe data dan referensi tabel dari setiap parameter
        4. Cari deskripsi dari text elements dan selection texts di bagian bawah kode
        5. Tentukan apakah parameter bersifat mandatory (OBLIGATORY)
        6. Identifikasi NO INTERVALS, NO-EXTENSION untuk SELECT-OPTIONS
        7. Cari default values dari INITIALIZATION atau DEFAULT VALUE
        8. Analisis validasi parameter di AT SELECTION-SCREEN

        Kembalikan JSON dengan format:
        {{
            "selection_parameters": [
                {{
                    "name": "nama_parameter_lengkap",
                    "type": "tipe_data_lengkap_dengan_referensi_tabel",
                    "description": "deskripsi lengkap dari text elements atau inferensi yang akurat",
                    "is_mandatory": true/false,
                    "is_select_option": true/false,
                    "has_no_intervals": true/false,
                    "default_value": "nilai_default_jika_ada",
                    "block_name": "nama_block_selection_screen",
                    "validation_logic": "logika_validasi_jika_ada"
                }}
            ]
        }}

        Kode ABAP:
        {raw_data.get('raw_code', '')}

        Komentar:
        {raw_data.get('raw_comments', '')}
        """
    
    def _create_error_handling_prompt(self, raw_data: Dict[str, Any]) -> str:
        """Error handling prompt"""
        return f"""
        Analisis kode ABAP secara komprehensif untuk semua skenario penanganan error.

        Kembalikan JSON:
        {{
            "error_scenarios": [
                {{
                    "error_description": "deskripsi kondisi error yang spesifik",
                    "resolution": "cara penyelesaian error yang detail",
                    "error_code": "kode error atau message ID jika ada",
                    "severity": "ERROR|WARNING|INFO",
                    "location": "di_bagian_mana_error_terjadi"
                }}
            ]
        }}

        Kode ABAP:
        {raw_data.get('raw_code', '')}
        """
    
    def _create_test_scenarios_prompt(self, raw_data: Dict[str, Any]) -> str:
        """Test scenarios prompt"""
        return f"""
        Berdasarkan analisis kode ABAP, buat skenario pengujian yang sangat komprehensif.

        Kembalikan JSON:
        {{
            "test_scenarios": [
                {{
                    "condition": "deskripsi kondisi pengujian yang detail",
                    "expected_result": "hasil yang diharapkan secara spesifik",
                    "test_data": "deskripsi data uji yang konkret",
                    "priority": "HIGH|MEDIUM|LOW",
                    "test_type": "POSITIVE|NEGATIVE|BOUNDARY|PERFORMANCE"
                }}
            ]
        }}

        Kode ABAP:
        {raw_data.get('raw_code', '')}
        """
    
    def _create_validation_rules_prompt(self, raw_data: Dict[str, Any]) -> str:
        """Validation rules prompt"""
        return f"""
        Analisis kode ABAP secara mendalam untuk semua aturan validasi data dan business rules.

        Kembalikan JSON:
        {{
            "validation_rules": [
                "deskripsi aturan validasi yang detail dan spesifik",
                "aturan validasi lainnya berdasarkan kode actual"
            ]
        }}

        Kode ABAP:
        {raw_data.get('raw_code', '')}
        """
    
    def _create_authorization_prompt(self, raw_data: Dict[str, Any]) -> str:
        """Authorization prompt"""
        return f"""
        Analisis kode ABAP secara mendalam untuk semua aspek otorisasi dan keamanan.

        Kembalikan JSON:
        {{
            "authorization_objects": ["daftar objek otorisasi yang ditemukan"],
            "user_roles": ["daftar peran/grup pengguna"],
            "authorization_logic": "deskripsi implementasi otorisasi yang detail"
        }}

        Kode ABAP:
        {raw_data.get('raw_code', '')}
        """
    
    async def _map_fixed_results_to_fsd(self, results: Dict[str, Dict[str, Any]]):
        """Map fixed comprehensive results to FSD document"""
        
        # Map basic info
        basic_info = results.get('basic_info', {})
        if basic_info:
            self.fsd_document.program_name = basic_info.get('program_name', '')
            self.fsd_document.report_description = basic_info.get('report_description', '')
            self.fsd_document.desain_report_description = basic_info.get('desain_report_description', '')
            self.fsd_document.user_requirements = basic_info.get('user_requirements', '')
            self.fsd_document.assumptions = basic_info.get('assumptions', [])
            self.fsd_document.transaction_code = basic_info.get('transaction_code', '')
            self.fsd_document.menu_path = basic_info.get('menu_path', 'N/A')
            
            # Add to related documents and version history
            ricefw_id = basic_info.get('ricefw_id', '')
            if ricefw_id:
                self.fsd_document.related_documents.append(f"RICEFW ID: {ricefw_id}")
            
            created_date = basic_info.get('created_date', '')
            created_by = basic_info.get('created_by', '')
            if created_date and created_by:
                self.fsd_document.version_history.append({
                    'version': '0.01',
                    'change': 'Initial draft',
                    'author': created_by,
                    'date': created_date
                })
            
            functional_contact = basic_info.get('functional_contact', '')
            if functional_contact:
                self.fsd_document.reviewers.append({
                    'role': 'Functional Lead',
                    'name': functional_contact
                })
        
        # Map selection parameters
        selection_screen = results.get('selection_screen', {})
        if selection_screen:
            for param_data in selection_screen.get('selection_parameters', []):
                self.fsd_document.selection_parameters.append(SelectionParameter(
                    name=param_data.get('name', ''),
                    type=param_data.get('type', ''),
                    description=param_data.get('description', ''),
                    is_mandatory=param_data.get('is_mandatory', False),
                    is_select_option=param_data.get('is_select_option', False),
                    has_no_intervals=param_data.get('has_no_intervals', False),
                    default_value=param_data.get('default_value', '')
                ))
        
        # Map COMPLETE field mappings
        field_mappings = results.get('complete_field_mappings', {})
        if field_mappings:
            for field_data in field_mappings.get('field_mappings', []):
                try:
                    processing_type = FieldProcessingType(field_data.get('processing_type', 'DIRECT'))
                except ValueError:
                    processing_type = FieldProcessingType.DIRECT
                
                self.fsd_document.field_mappings.append(FieldMapping(
                    display_name=field_data.get('display_name', ''),
                    technical_field=field_data.get('technical_field', ''),
                    source_table=field_data.get('source_table', ''),
                    processing_logic=field_data.get('processing_logic', ''),
                    processing_type=processing_type,
                    join_condition=field_data.get('join_condition', ''),
                    where_condition=field_data.get('where_condition', '')
                ))
        
        # Map COMPLETE valid dataset rules
        valid_datasets = results.get('complete_valid_datasets', {})
        if valid_datasets:
            for dataset_rule in valid_datasets.get('valid_dataset_rules', []):
                self.fsd_document.valid_dataset_rules.append(DataConditionRow(
                    data=dataset_rule.get('data', ''),
                    condition=dataset_rule.get('condition', '')
                ))
        
        # Map COMPLETE lookup forms with proper categorization
        lookup_forms = results.get('complete_lookup_forms', {})
        if lookup_forms:
            # Map company code lookup to country_info
            for company_rule in lookup_forms.get('company_code_lookup', []):
                condition_text = company_rule.get('condition', '')
                target_fields = company_rule.get('target_fields', ['BUTXT'])
                self.fsd_document.country_info.append(DataConditionRow(
                    data=' & '.join(target_fields),
                    condition=condition_text
                ))
            
            # Map business area lookup to currency_t500c
            for ba_rule in lookup_forms.get('business_area_lookup', []):
                condition_text = ba_rule.get('condition', '')
                target_fields = ba_rule.get('target_fields', ['GTEXT'])
                self.fsd_document.currency_t500c.append(DataConditionRow(
                    data=' & '.join(target_fields),
                    condition=condition_text
                ))
            
            # Map wage type lookup to currency_t001
            for wt_rule in lookup_forms.get('wage_type_lookup', []):
                condition_text = wt_rule.get('condition', '')
                target_fields = wt_rule.get('target_fields', ['LGTXT'])
                self.fsd_document.currency_t001.append(DataConditionRow(
                    data=' & '.join(target_fields),
                    condition=condition_text
                ))
                
            # Add payroll reason lookup if exists (create additional section if needed)
            for pr_rule in lookup_forms.get('payroll_reason_lookup', []):
                condition_text = pr_rule.get('condition', '')
                target_fields = pr_rule.get('target_fields', ['OCRTX'])
                # Add to currency_t001 or create separate section
                self.fsd_document.currency_t001.append(DataConditionRow(
                    data=' & '.join(target_fields),
                    condition=condition_text
                ))
        
        # Map error scenarios, test scenarios, validation rules, authorization (same as before)
        error_handling = results.get('error_handling', {})
        if error_handling:
            for error_data in error_handling.get('error_scenarios', []):
                self.fsd_document.error_scenarios.append(ErrorScenario(
                    error_description=error_data.get('error_description', ''),
                    resolution=error_data.get('resolution', ''),
                    error_code=error_data.get('error_code', ''),
                    severity=error_data.get('severity', 'ERROR')
                ))
        
        test_scenarios = results.get('test_scenarios', {})
        if test_scenarios:
            for test_data in test_scenarios.get('test_scenarios', []):
                self.fsd_document.test_scenarios.append(TestScenario(
                    condition=test_data.get('condition', ''),
                    expected_result=test_data.get('expected_result', ''),
                    test_data=test_data.get('test_data', ''),
                    priority=test_data.get('priority', 'HIGH')
                ))
        
        validation_rules = results.get('validation_rules', {})
        if validation_rules:
            self.fsd_document.validation_rules.extend(validation_rules.get('validation_rules', []))
        
        authorization = results.get('authorization', {})
        if authorization:
            self.fsd_document.authorization_objects.extend(authorization.get('authorization_objects', []))
            self.fsd_document.user_roles.extend(authorization.get('user_roles', []))

# ================================
# ENHANCED INTELLIGENT FSD GENERATOR
# ================================

class EnhancedIntelligentFSDGenerator:
    """Main orchestrator class for intelligent FSD generation with Word template support"""
    
    def __init__(self, config_file: str = None):
        self.config = ConfigManager(config_file)
        self.config.validate_required()
        self.mapper = IntelligentFSDMapper(self.config)
        self.output_generator = EnhancedOutputGenerator(self.config)
    
    async def process_file(self, html_file_path: str, template_path: str = None, 
                          custom_output_dir: str = None) -> Dict[str, Any]:
        """Process a single HTML file and generate FSD outputs including Word document"""
        logger.info(f"Processing file: {html_file_path}")
        
        if not os.path.exists(html_file_path):
            raise FileNotFoundError(f"HTML file not found: {html_file_path}")
        
        # Use provided template path or default
        if not template_path:
            # template_paths = "/Users/wilbert.limson/python_project/PLN-Genie/Template/FSD/PLN_SI SSoT_(DAPI ID)_(Module Name)_Functional Specification Design (FSD)_v100_ID.docx"
            template_path = r"C:\Users\wahyu.perwira\Documents\Project\poc\SAP-AUTOMATE-FD-TD\backend\templates\Template_PLN_SI SSoT_(DAPI ID)_(Module Name)_Functional Specification Design (FSD)_v100_ID.docx"
        
        # Generate FSD document
        fsd_document = await self.mapper.analyze_and_map(html_file_path)
        
        # Generate outputs
        base_filename = Path(html_file_path).stem
        if custom_output_dir:
            original_output_dir = self.output_generator.output_dir
            self.output_generator.output_dir = Path(custom_output_dir)
            self.output_generator.output_dir.mkdir(parents=True, exist_ok=True)
        
        # output_files = self.output_generator.generate_all_outputs(fsd_document, base_filename, template_path)
        output_files = self.output_generator.generate_all_outputs(
            fsd_document,
            base_filename,
            template_path
        )

        # Generate final Word document using md_to_docs_converter
        try:
            markdown_path = output_files.get('markdown')
            if markdown_path:
                converter = TitlePageGenerator(template_path)
                doc_result = converter.generate_complete_document(
                    markdown_path,
                    str(self.output_generator.output_dir)
                )
                if doc_result.get('word_path'):
                    output_files['final_docx'] = doc_result['word_path']
        except Exception as exc:
            logger.error(f"Failed to convert markdown to final DOCX: {exc}")
        
        if custom_output_dir:
            self.output_generator.output_dir = original_output_dir
        
        # Compile results
        results = {
            'input_file': html_file_path,
            'template_file': template_path,
            'fsd_document': dataclass_to_dict(fsd_document),
            # 'output_files': asdict(fsd_document),
            'output_files': output_files,
            'analysis_summary': {
                'program_name': fsd_document.program_name,
                'description': fsd_document.report_description,
                'desain_description': fsd_document.desain_report_description,
                'selection_parameters_count': len(fsd_document.selection_parameters),
                'field_mappings_count': len(fsd_document.field_mappings),
                'error_scenarios_count': len(fsd_document.error_scenarios),
                'test_scenarios_count': len(fsd_document.test_scenarios),
                'validation_rules_count': len(fsd_document.validation_rules),
                'authorization_objects_count': len(fsd_document.authorization_objects)
            }
        }
        
        logger.info(f"Successfully processed {html_file_path}")
        return results
    
    async def process_multiple_files(self, html_files: List[str], template_path: str = None, 
                                   output_dir: str = None) -> Dict[str, Any]:
        """Process multiple HTML files"""
        logger.info(f"Processing {len(html_files)} files")
        
        results = {}
        successful = 0
        failed = 0
        
        for html_file in html_files:
            try:
                file_results = await self.process_file(html_file, template_path, output_dir)
                results[html_file] = file_results
                successful += 1
                logger.info(f"✓ Successfully processed: {os.path.basename(html_file)}")
            except Exception as e:
                logger.error(f"✗ Failed to process {os.path.basename(html_file)}: {e}")
                results[html_file] = {'error': str(e)}
                failed += 1
        
        summary = {
            'total_files': len(html_files),
            'successful': successful,
            'failed': failed,
            'results': results
        }
        
        logger.info(f"Batch processing complete: {successful} successful, {failed} failed")
        return summary

# ================================
# FASTAPI INTEGRATION
# ================================

# FastAPI Pydantic models
class ProcessingRequest(BaseModel):
    file_paths: List[str]
    template_path: Optional[str] = None
    output_dir: Optional[str] = None
    config: Optional[Dict[str, Any]] = None

class ConfigurationRequest(BaseModel):
    gemini_api_key: str
    project_name: Optional[str] = "System Integrator for MIS Towards SSoT"
    max_tokens: Optional[int] = 4096
    temperature: Optional[float] = 0.1
    gemini_api_url: Optional[str] = "https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-pro-latest:generateContent"
    requirement_list_excel: Optional[str] = "/Users/wahyu.perwira/Documents/Project/poc/SAP-AUTOMATE-FD-TD/backend/output/database/Requirement-List.xlsx"

class ProcessingStatus(BaseModel):
    job_id: str
    status: str  # "pending", "processing", "completed", "failed"
    progress: int
    message: str
    results: Optional[Dict[str, Any]] = None
    error: Optional[str] = None

class FileDeleteRequest(BaseModel):
    path: str

# FastAPI app initialization
app = FastAPI(
    title="Accenture SAP FSD Document Processor", 
    version="2.0.0",
    description="AI-Powered SAP ABAP Functional Specification Design Generator"
)

# Configure CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "http://localhost:3000", "http://localhost:8501"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Directory setup
UPLOAD_DIR = "uploads"
OUTPUT_DIR = "/Users/wahyu.perwira/Documents/Project/poc/SAP-AUTOMATE-FD-TD/backend/output/output"
TEMPLATE_DIR = "templates"
CONFIG_DIR = "config"

for directory in [UPLOAD_DIR, OUTPUT_DIR, TEMPLATE_DIR, CONFIG_DIR]:
    os.makedirs(directory, exist_ok=True)

# Global variables
stored_files = []
processing_jobs = {}
fsd_generator = None

# Initialize FSD Generator
async def initialize_fsd_generator(config_data: Dict[str, Any] = None):
    """Initialize the FSD Generator with configuration"""
    global fsd_generator
    
    try:
        # Create temporary config file if config_data provided
        if config_data:
            config_file = os.path.join(CONFIG_DIR, "temp_config.json")
            
            # Set environment variables
            for key, value in config_data.items():
                os.environ[key.upper()] = str(value)
            
            # Save config to file
            with open(config_file, 'w', encoding='utf-8') as f:
                json.dump(config_data, f, indent=2)
        else:
            config_file = None
        
        # Initialize the generator
        fsd_generator = EnhancedIntelligentFSDGenerator(config_file)
        logger.info("FSD Generator initialized successfully")
            
    except Exception as e:
        logger.error(f"Error initializing FSD Generator: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to initialize FSD Generator: {str(e)}")

@app.on_event("startup")
async def startup_event():
    """Initialize the application on startup"""
    logger.info("Starting Accenture SAP FSD Document Processor")
    
    # Try to initialize with default config
    try:
        await initialize_fsd_generator()
    except Exception as e:
        logger.warning(f"Could not initialize FSD Generator on startup: {e}")

@app.get("/")
async def root():
    return {
        "message": "Accenture SAP FSD Document Processor API",
        "status": "running",
        "version": "2.0.0",
        "fsd_generator_available": fsd_generator is not None,
        "features": [
            "AI-Powered ABAP Analysis",
            "Multi-format Document Generation",
            "Batch Processing",
            "Word Template Support"
        ]
    }

@app.post("/api/configure")
async def configure_fsd_generator(config: ConfigurationRequest):
    """Configure the FSD Generator with API keys and settings"""
    try:
        config_data = {
            "GEMINI_API_KEY": config.gemini_api_key,
            "gemini_api_url": config.gemini_api_url,
            "project_name": config.project_name,
            "max_tokens": config.max_tokens,
            "temperature": config.temperature,
            "requirement_list_excel": config.requirement_list_excel,
            "default_output_dir": OUTPUT_DIR,
            "template_dir": TEMPLATE_DIR
        }
        
        await initialize_fsd_generator(config_data)
        
        return JSONResponse(content={
            "success": True,
            "message": "FSD Generator configured successfully",
            "config": {k: v for k, v in config_data.items() if k != "GEMINI_API_KEY"}
        })
        
    except Exception as e:
        logger.error(f"Error configuring FSD Generator: {e}")
        raise HTTPException(status_code=500, detail=f"Configuration error: {str(e)}")

@app.post("/api/store-file")
async def store_file_locally(file: UploadFile = File(...), originalName: str = None):
    """Store uploaded HTML file in local file system"""
    try:
        # Validate file type
        if not file.filename.endswith(('.html', '.htm')):
            raise HTTPException(status_code=400, detail="Only HTML files are allowed")
        
        # Generate unique filename with timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_name = os.path.splitext(originalName or file.filename)[0]
        filename = f"{base_name}_{timestamp}.html"
        file_path = os.path.join(UPLOAD_DIR, filename)
        
        # Save file to local file system
        with open(file_path, "wb") as buffer:
            content = await file.read()
            buffer.write(content)
        
        # Store file info for tracking
        file_info = {
            "id": str(uuid.uuid4()),
            "name": originalName or file.filename,
            "filename": filename,
            "path": file_path,
            "size": len(content),
            "type": file.content_type,
            "uploadedAt": datetime.now().isoformat(),
            "status": "stored"
        }
        stored_files.append(file_info)
        
        logger.info(f"File stored at: {file_path}")
        
        return JSONResponse(content={
            "success": True,
            "file_info": file_info,
            "message": f"File stored successfully"
        })
        
    except Exception as e:
        logger.error(f"Error storing file: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error storing file: {str(e)}")

@app.post("/api/upload-template")
async def upload_template(template: UploadFile = File(...)):
    """Upload Word template file"""
    try:
        # Validate file type
        if not template.filename.endswith('.docx'):
            raise HTTPException(status_code=400, detail="Only .docx template files are allowed")
        
        # Save template file
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"template_{timestamp}.docx"
        template_path = os.path.join(TEMPLATE_DIR, filename)
        
        with open(template_path, "wb") as buffer:
            content = await template.read()
            buffer.write(content)
        
        logger.info(f"Template uploaded: {template_path}")
        
        return JSONResponse(content={
            "success": True,
            "template_path": template_path,
            "filename": filename,
            "message": "Template uploaded successfully"
        })
        
    except Exception as e:
        logger.error(f"Error uploading template: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error uploading template: {str(e)}")

@app.get("/api/list-files")
async def list_stored_files():
    """Get list of all stored files"""
    try:
        # Refresh stored files list from directory
        actual_files = []
        if os.path.exists(UPLOAD_DIR):
            for filename in os.listdir(UPLOAD_DIR):
                if filename.endswith(('.html', '.htm')):
                    file_path = os.path.join(UPLOAD_DIR, filename)
                    stat = os.stat(file_path)
                    actual_files.append({
                        "id": str(uuid.uuid4()),
                        "name": filename,
                        "filename": filename,
                        "path": file_path,
                        "size": stat.st_size,
                        "type": "text/html",
                        "uploadedAt": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                        "status": "stored"
                    })
        
        return JSONResponse(content={
            "success": True,
            "files": actual_files,
            "count": len(actual_files)
        })
    except Exception as e:
        logger.error(f"Error listing files: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error listing files: {str(e)}")

@app.post("/api/process-files")
async def process_files(
    background_tasks: BackgroundTasks,
    request: ProcessingRequest
):
    """Process HTML files using the FSD Generator"""
    try:
        if not fsd_generator:
            raise HTTPException(
                status_code=503, 
                detail="FSD Generator not available. Please configure first."
            )
        
        # Validate file paths
        valid_files = []
        for file_path in request.file_paths:
            if os.path.exists(file_path) and file_path.endswith(('.html', '.htm')):
                valid_files.append(file_path)
            else:
                logger.warning(f"Invalid file path: {file_path}")
        
        if not valid_files:
            raise HTTPException(status_code=400, detail="No valid HTML files found")
        
        # Generate job ID
        job_id = str(uuid.uuid4())
        
        # Initialize job status
        processing_jobs[job_id] = ProcessingStatus(
            job_id=job_id,
            status="pending",
            progress=0,
            message="Job queued for processing"
        )
        
        # Start background processing
        background_tasks.add_task(
            process_files_background,
            job_id,
            valid_files,
            request.template_path,
            request.output_dir or OUTPUT_DIR
        )
        
        return JSONResponse(content={
            "success": True,
            "job_id": job_id,
            "files_count": len(valid_files),
            "message": "Processing started in background"
        })
        
    except Exception as e:
        logger.error(f"Error starting file processing: {e}")
        raise HTTPException(status_code=500, detail=f"Error starting processing: {str(e)}")

# Fix 5: Update the background processing function to store serializable results
async def process_files_background(
    job_id: str,
    file_paths: List[str],
    template_path: Optional[str],
    output_dir: str
):
    """Background task for processing files with the FSD Generator"""
    try:
        # Update job status
        processing_jobs[job_id].status = "processing"
        processing_jobs[job_id].message = "Initializing AI analysis..."
        processing_jobs[job_id].progress = 10
        
        results = {}
        
        if len(file_paths) == 1:
            # Single file processing
            processing_jobs[job_id].message = f"Processing {os.path.basename(file_paths[0])}"
            processing_jobs[job_id].progress = 30
            
            result = await fsd_generator.process_file(
                file_paths[0], 
                template_path, 
                output_dir
            )
            results = {'single': result}
            
        else:
            # Batch processing
            processing_jobs[job_id].message = f"Processing {len(file_paths)} files"
            processing_jobs[job_id].progress = 30
            
            result = await fsd_generator.process_multiple_files(
                file_paths, 
                template_path, 
                output_dir
            )
            results = {'batch': result}
        
        # **FIX: Convert results to JSON-serializable format before storing**
        serializable_results = dataclass_to_dict(results)
        
        # Update job status to completed
        processing_jobs[job_id].status = "completed"
        processing_jobs[job_id].progress = 100
        processing_jobs[job_id].message = "Processing completed successfully"
        processing_jobs[job_id].results = serializable_results
        
        logger.info(f"Job {job_id} completed successfully")
        
    except Exception as e:
        # Update job status to failed
        processing_jobs[job_id].status = "failed"
        processing_jobs[job_id].progress = 0
        processing_jobs[job_id].message = "Processing failed"
        processing_jobs[job_id].error = str(e)
        
        logger.error(f"Job {job_id} failed: {e}")

# Fix 3: Update get_job_status function to handle serialization
@app.get("/api/job-status/{job_id}")
async def get_job_status(job_id: str):
    """Get the status of a processing job"""
    try:
        if job_id not in processing_jobs:
            raise HTTPException(status_code=404, detail="Job not found")
        
        job_status = processing_jobs[job_id]
        
        # **FIX: Convert results to JSON-serializable format if they exist**
        serializable_results = None
        if job_status.results:
            serializable_results = dataclass_to_dict(job_status.results)
        
        return JSONResponse(content={
            "success": True,
            "job_id": job_id,
            "status": job_status.status,
            "progress": job_status.progress,
            "message": job_status.message,
            "results": serializable_results,
            "error": job_status.error
        })
        
    except Exception as e:
        logger.error(f"Error getting job status: {e}")
        raise HTTPException(status_code=500, detail=f"Error getting job status: {str(e)}")

# Fix 2: Update get_job_results function to handle serialization
@app.get("/api/job-results/{job_id}")
async def get_job_results(job_id: str):
    """Get detailed results of a completed processing job"""
    try:
        if job_id not in processing_jobs:
            raise HTTPException(status_code=404, detail="Job not found")
        
        job_status = processing_jobs[job_id]
        
        if job_status.status != "completed":
            raise HTTPException(status_code=400, detail="Job not completed yet")
        
        if not job_status.results:
            raise HTTPException(status_code=404, detail="No results available")
        
        # **FIX: Convert results to JSON-serializable format**
        serializable_results = dataclass_to_dict(job_status.results)
        
        # Process results for API response
        processed_results = {}
        
        if 'single' in serializable_results:
            result = serializable_results['single']
            processed_results = {
                'type': 'single',
                'input_file': result.get('input_file'),
                'template_file': result.get('template_file'),
                'output_files': result.get('output_files'),
                'analysis_summary': result.get('analysis_summary')
            }
        elif 'batch' in serializable_results:
            batch_result = serializable_results['batch']
            processed_results = {
                'type': 'batch',
                'total_files': batch_result.get('total_files'),
                'successful': batch_result.get('successful'),
                'failed': batch_result.get('failed'),
                'results_summary': []
            }
            
            # Add summary for each file
            for file_path, result in batch_result.get('results', {}).items():
                if 'error' not in result:
                    processed_results['results_summary'].append({
                        'file': os.path.basename(file_path),
                        'status': 'success',
                        'program_name': result.get('analysis_summary', {}).get('program_name'),
                        'output_files': result.get('output_files')
                    })
                else:
                    processed_results['results_summary'].append({
                        'file': os.path.basename(file_path),
                        'status': 'failed',
                        'error': result.get('error')
                    })
        
        return JSONResponse(content={
            "success": True,
            "job_id": job_id,
            "results": processed_results
        })
        
    except Exception as e:
        logger.error(f"Error getting job results: {e}")
        raise HTTPException(status_code=500, detail=f"Error getting job results: {str(e)}")

@app.get("/api/download-file")
async def download_generated_file(file_path: str):
    """Download a generated file"""
    try:
        if not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail="File not found")
        
        # Determine media type based on file extension
        if file_path.endswith('.docx'):
            media_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        elif file_path.endswith('.json'):
            media_type = 'application/json'
        elif file_path.endswith('.md'):
            media_type = 'text/markdown'
        else:
            media_type = 'text/plain'
        
        filename = os.path.basename(file_path)
        
        return FileResponse(
            path=file_path,
            media_type=media_type,
            filename=filename
        )
        
    except Exception as e:
        logger.error(f"Error downloading file: {e}")
        raise HTTPException(status_code=500, detail=f"Error downloading file: {str(e)}")

@app.get("/api/file-content")
async def get_file_content(file_path: str, encoding: str = "utf-8"):
    """Get the content of a text file"""
    try:
        if not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail="File not found")
        
        # Only allow text files for content viewing
        allowed_extensions = ['.md', '.txt', '.json', '.html', '.htm']
        if not any(file_path.endswith(ext) for ext in allowed_extensions):
            raise HTTPException(status_code=400, detail="File type not supported for content viewing")
        
        with open(file_path, 'r', encoding=encoding) as f:
            content = f.read()
        
        return JSONResponse(content={
            "success": True,
            "file_path": file_path,
            "content": content,
            "size": len(content)
        })
        
    except Exception as e:
        logger.error(f"Error reading file content: {e}")
        raise HTTPException(status_code=500, detail=f"Error reading file content: {str(e)}")

@app.delete("/api/delete-file")
async def delete_stored_file(request_data: FileDeleteRequest):
    """Delete a stored file from local file system"""
    try:
        file_path = request_data.path
        if not file_path:
            raise HTTPException(status_code=400, detail="File path is required")
        
        if os.path.exists(file_path):
            os.remove(file_path)
            logger.info(f"File deleted: {file_path}")
            
            # Remove from tracking list
            global stored_files
            stored_files = [f for f in stored_files if f["path"] != file_path]
            
            return JSONResponse(content={
                "success": True,
                "message": f"File deleted successfully"
            })
        else:
            raise HTTPException(status_code=404, detail="File not found")
            
    except Exception as e:
        logger.error(f"Error deleting file: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error deleting file: {str(e)}")

@app.delete("/api/clear-files")
async def clear_all_files():
    """Clear all stored files from local file system"""
    try:
        # Clear uploads
        if os.path.exists(UPLOAD_DIR):
            for filename in os.listdir(UPLOAD_DIR):
                file_path = os.path.join(UPLOAD_DIR, filename)
                if os.path.isfile(file_path):
                    os.remove(file_path)
        
        # Clear outputs (optional)
        if os.path.exists(OUTPUT_DIR):
            for filename in os.listdir(OUTPUT_DIR):
                file_path = os.path.join(OUTPUT_DIR, filename)
                if os.path.isfile(file_path):
                    os.remove(file_path)
        
        global stored_files, processing_jobs
        stored_files = []
        processing_jobs = {}
        
        logger.info("All files and jobs cleared")
        
        return JSONResponse(content={
            "success": True,
            "message": "All files and jobs cleared successfully"
        })
        
    except Exception as e:
        logger.error(f"Error clearing files: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error clearing files: {str(e)}")

@app.get("/api/system-info")
async def get_system_info():
    """Get system information and status"""
    try:
        # Check directory status
        directories_info = {}
        for dir_name, dir_path in [
            ("uploads", UPLOAD_DIR),
            ("output", OUTPUT_DIR),
            ("templates", TEMPLATE_DIR),
            ("config", CONFIG_DIR)
        ]:
            directories_info[dir_name] = {
                "path": dir_path,
                "exists": os.path.exists(dir_path),
                "file_count": len(os.listdir(dir_path)) if os.path.exists(dir_path) else 0
            }
        
        # Check FSD Generator status
        generator_info = {
            "generator_initialized": fsd_generator is not None,
            "active_jobs": len(processing_jobs),
            "stored_files": len(stored_files)
        }
        
        return JSONResponse(content={
            "success": True,
            "system_info": {
                "directories": directories_info,
                "generator": generator_info,
                "version": "2.0.0",
                "timestamp": datetime.now().isoformat()
            }
        })
        
    except Exception as e:
        logger.error(f"Error getting system info: {e}")
        raise HTTPException(status_code=500, detail=f"Error getting system info: {str(e)}")
    
@app.post("/api/process-html")
async def process_html_file(html_file: UploadFile = File(...)):
    """Process HTML file directly from upload without storing first."""
    try:
        logger.info(f"📥 Received file upload request: {html_file.filename}")

        # Validate file type
        if not html_file.filename.endswith(('.html', '.htm')):
            logger.warning(f"❌ Invalid file type: {html_file.filename}")
            raise HTTPException(status_code=400, detail="Only HTML files are allowed")
        
        # Read file content directly into memory
        logger.info("📄 Reading file content from upload...")
        file_bytes = await html_file.read()
        logger.info(f"✅ Read {len(file_bytes)} bytes from file: {html_file.filename}")

        # Save temporarily if fsd_generator requires a file path
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        temp_filename = f"temp_{timestamp}.html"
        temp_path = os.path.join(OUTPUT_DIR, temp_filename)

        logger.info(f"💾 Saving temp file to: {temp_path}")
        with open(temp_path, "wb") as f:
            f.write(file_bytes)
        logger.info("✅ Temp file saved successfully.")

        # Initialize generator if needed
        if not fsd_generator:
            logger.info("⚙ Initializing FSD Generator...")
            try:
                await initialize_fsd_generator()
                logger.info("✅ FSD Generator initialized.")
            except Exception as init_err:
                logger.error(f"❌ Failed to initialize FSD Generator: {init_err}")
                raise HTTPException(
                    status_code=503,
                    detail="FSD Generator not available. Please configure first."
                )

        # Process file directly
        logger.info(f"🚀 Processing file: {temp_path}")
        result = await fsd_generator.process_file(temp_path, None, OUTPUT_DIR)
        logger.info("✅ File processed successfully by FSD Generator.")

        # Read markdown output (if exists)
        markdown_path = result.get('output_files', {}).get('markdown')
        markdown_content = ""
        if markdown_path and os.path.exists(markdown_path):
            logger.info(f"📄 Reading markdown output from: {markdown_path}")
            with open(markdown_path, 'r', encoding='utf-8') as md_file:
                markdown_content = md_file.read()
            logger.info(f"✅ Markdown file read successfully ({len(markdown_content)} chars).")
        else:
            logger.info("ℹ No markdown output found.")

        # Convert to JSON-safe format
        logger.info("🔄 Converting result to JSON-serializable format...")
        serializable_result = dataclass_to_dict(result)
        logger.info("✅ Conversion complete.")

        # Delete temp file after processing
        try:
            os.remove(temp_path)
            logger.info(f"🗑 Temp file deleted: {temp_path}")
        except OSError as del_err:
            logger.warning(f"⚠ Could not delete temp file: {del_err}")

        return JSONResponse(content={
            "success": True,
            "filename": html_file.filename,
            "processed_data": serializable_result,
            "fsd_analysis": serializable_result.get('analysis_summary'),
            "output_files": serializable_result.get('output_files'),
            "markdown_content": markdown_content,
            "message": "HTML file processed successfully"
        })

    except Exception as e:
        logger.error(f"❌ Error processing HTML file: {str(e)}")
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Error processing file: {str(e)}")

@app.post("/api/process-html-v1")
async def process_html_file_v1(html_file: UploadFile = File(...)):
    """Process single HTML file - Legacy endpoint for backward compatibility"""
    try:
        # Store the file first
        store_result = await store_file_locally(html_file)
        
        if not store_result:
            raise HTTPException(status_code=500, detail="Failed to store file")
        
        # Extract file info from store result
        response_data = json.loads(store_result.body.decode())
        file_info = response_data["file_info"]
        file_path = file_info["path"]
        
        # Process immediately (not in background for legacy compatibility)
        if not fsd_generator:
            try:
                await initialize_fsd_generator()
            except Exception as init_err:
                logger.error(f"Failed to initialize FSD Generator: {init_err}")
                raise HTTPException(
                    status_code=503,
                    detail="FSD Generator not available. Please configure first."
                )
        
        # Call the main processing logic
        result = await fsd_generator.process_file(file_path, None, OUTPUT_DIR)

        markdown_path = result.get('output_files', {}).get('markdown')
        markdown_content = ""
        if markdown_path and os.path.exists(markdown_path):
            with open(markdown_path, 'r', encoding='utf-8') as md_file:
                markdown_content = md_file.read()
        
        # **FIX: Convert all data to JSON-serializable format before returning**
        serializable_result = dataclass_to_dict(result)
        
        return JSONResponse(content={
            "success": True,
            "file_path": file_path,
            "filename": html_file.filename,
            "processed_data": serializable_result,
            "fsd_analysis": serializable_result.get('analysis_summary'),
            "output_files": serializable_result.get('output_files'),
            "markdown_content": markdown_content,
            "message": "HTML file processed successfully"
        })
        
    except Exception as e:
        logger.error(f"Error processing HTML file: {str(e)}")
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Error processing file: {str(e)}")

@app.get("/api/health")
async def health_check():
    """Health check endpoint"""
    return {
        "status": "healthy",
        "timestamp": datetime.now().isoformat(),
        "fsd_generator_available": fsd_generator is not None,
        "version": "2.0.0"
    }

# ================================
# ADVANCED WORD TEMPLATE PROCESSOR
# ================================

class AdvancedWordTemplateProcessor(WordTemplateProcessor):
    """Advanced Word template processor with better formatting preservation"""
    
    def __init__(self, template_path: str):
        super().__init__(template_path)
        self.placeholder_map = {}
        self.style_preservation = True
    
    def detect_placeholders(self):
        """Detect placeholders in the template"""
        if self.template_doc:
            # For DocxTemplate, placeholders are Jinja2 style
            content = self.template_doc.docx.main_document_part.element.xml
            placeholders = re.findall(r'\{\{([^}]+)\}\}', content)
            self.placeholder_map = {f"{{{{{p}}}}}": p for p in placeholders}
        else:
            # For regular Document, look for custom placeholder patterns
            for paragraph in self.document.paragraphs:
                placeholders = re.findall(r'\[([A-Z_]+)\]', paragraph.text)
                for p in placeholders:
                    self.placeholder_map[f"[{p}]"] = p.lower()
    
    def preserve_formatting(self, paragraph, new_text: str):
        """Preserve formatting while replacing text"""
        if not paragraph.runs:
            paragraph.text = new_text
            return
        
        # Keep the formatting of the first run
        first_run = paragraph.runs[0]
        font = first_run.font
        
        # Clear all runs
        for run in paragraph.runs:
            run.clear()
        
        # Add new text with preserved formatting
        new_run = paragraph.runs[0]
        new_run.text = new_text
        new_run.font.name = font.name
        new_run.font.size = font.size
        new_run.font.bold = font.bold
        new_run.font.italic = font.italic
    
    def fill_template_with_enhanced_data(self, fsd_document: FSDDocument, output_path: str):
        """Fill template with enhanced data mapping"""
        # Prepare comprehensive context
        context = {
            # Basic information
            'program_name': fsd_document.program_name,
            'report_description': fsd_document.report_description,
            'current_date': datetime.now().strftime('%d %B %Y'),
            'project_name': fsd_document.project_name,
            'transaction_code': fsd_document.transaction_code,
            'menu_path': fsd_document.menu_path,
            
            # Document information
            'document_location': fsd_document.document_location,
            'related_documents': fsd_document.related_documents,
            'reviewers': fsd_document.reviewers,
            'version_history': fsd_document.version_history,
            
            # Requirements
            'user_requirements': fsd_document.user_requirements,
            'assumptions': fsd_document.assumptions,
            
            # Design elements
            'selection_parameters': [asdict(p) for p in fsd_document.selection_parameters],
            'field_mappings': [asdict(f) for f in fsd_document.field_mappings],
            'validation_rules': fsd_document.validation_rules,
            
            # Error and testing
            'error_scenarios': [asdict(e) for e in fsd_document.error_scenarios],
            'test_scenarios': [asdict(t) for t in fsd_document.test_scenarios],
            
            # Authorization
            'authorization_objects': fsd_document.authorization_objects,
            'user_roles': fsd_document.user_roles,
            
            # Statistics
            'stats': {
                'selection_parameters_count': len(fsd_document.selection_parameters),
                'field_mappings_count': len(fsd_document.field_mappings),
                'error_scenarios_count': len(fsd_document.error_scenarios),
                'test_scenarios_count': len(fsd_document.test_scenarios)
            }
        }
        
        if self.template_doc:
            try:
                self.template_doc.render(context)
                self.template_doc.save(output_path)
                logger.info(f"Successfully filled template using DocxTemplate: {output_path}")
            except Exception as e:
                logger.error(f"DocxTemplate failed: {e}")
                self._fallback_manual_fill(context, output_path)
        else:
            self._fallback_manual_fill(context, output_path)
    
    def _fallback_manual_fill(self, context: Dict[str, Any], output_path: str):
        """Fallback manual filling method"""
        try:
            if not self.document:
                self.document = DocxDocument(self.template_path)
            
            # Replace simple placeholders
            simple_replacements = {
                '[PROGRAM_NAME]': context.get('program_name', ''),
                '[DESCRIPTION]': context.get('report_description', ''),
                '[CURRENT_DATE]': context.get('current_date', ''),
                '[PROJECT_NAME]': context.get('project_name', ''),
                '[TRANSACTION_CODE]': context.get('transaction_code', ''),
                '[USER_REQUIREMENTS]': context.get('user_requirements', ''),
                '{{program_name}}': context.get('program_name', ''),
                '{{description}}': context.get('report_description', ''),
                '{{current_date}}': context.get('current_date', ''),
                '{{project_name}}': context.get('project_name', '')
            }
            
            # Replace in paragraphs
            for paragraph in self.document.paragraphs:
                for placeholder, value in simple_replacements.items():
                    if placeholder in paragraph.text:
                        if self.style_preservation:
                            self.preserve_formatting(paragraph, paragraph.text.replace(placeholder, value))
                        else:
                            paragraph.text = paragraph.text.replace(placeholder, value)
            
            # Replace in tables
            for table in self.document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for placeholder, value in simple_replacements.items():
                                if placeholder in paragraph.text:
                                    if self.style_preservation:
                                        self.preserve_formatting(paragraph, paragraph.text.replace(placeholder, value))
                                    else:
                                        paragraph.text = paragraph.text.replace(placeholder, value)
            
            # Add dynamic content sections
            self._add_dynamic_sections(context)
            
            self.document.save(output_path)
            logger.info(f"Successfully filled template using manual method: {output_path}")
            
        except Exception as e:
            logger.error(f"Manual template fill failed: {e}")
            raise
    
    def _add_dynamic_sections(self, context: Dict[str, Any]):
        """Add dynamic sections to the document"""
        # Find insertion points and add content
        # This is a simplified version - in practice, you'd need more sophisticated logic
        # to find the right insertion points in the template
        
        # Add selection parameters table if found
        selection_params = context.get('selection_parameters', [])
        if selection_params:
            self._add_selection_parameters_table(selection_params)
        
        # Add field mappings table if found
        field_mappings = context.get('field_mappings', [])
        if field_mappings:
            self._add_field_mappings_table(field_mappings)
    
    def _add_selection_parameters_table(self, parameters: List[Dict]):
        """Add selection parameters table to document"""
        # Find a good insertion point (after "Selection Screen" heading)
        for paragraph in self.document.paragraphs:
            if "Selection Screen" in paragraph.text or "4.1" in paragraph.text:
                # Insert table after this paragraph
                break
        
        # Create table
        if parameters:
            table = self.document.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            
            # Header row
            header_cells = table.rows[0].cells
            headers = ['Parameter', 'Type', 'Description', 'Mandatory', 'Select-Option', 'No Intervals']
            for i, header in enumerate(headers):
                header_cells[i].text = header
                # Make header bold
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # Data rows
            for param in parameters:
                row_cells = table.add_row().cells
                row_cells[0].text = param.get('name', '')
                row_cells[1].text = param.get('type', '')
                row_cells[2].text = param.get('description', '')
                row_cells[3].text = 'Yes' if param.get('is_mandatory') else 'No'
                row_cells[4].text = 'Yes' if param.get('is_select_option') else 'No'
                row_cells[5].text = 'Yes' if param.get('has_no_intervals') else 'No'
    
    def _add_field_mappings_table(self, field_mappings: List[Dict]):
        """Add field mappings table to document"""
        # Similar logic for field mappings
        if field_mappings:
            table = self.document.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            
            # Header row
            header_cells = table.rows[0].cells
            headers = ['Field Name', 'Technical Field', 'Source Table', 'Processing Logic', 'Processing Type']
            for i, header in enumerate(headers):
                header_cells[i].text = header
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # Data rows
            for field in field_mappings:
                row_cells = table.add_row().cells
                row_cells[0].text = field.get('display_name', '')
                row_cells[1].text = field.get('technical_field', '')
                row_cells[2].text = field.get('source_table', '')
                row_cells[3].text = field.get('processing_logic', '')
                row_cells[4].text = field.get('processing_type', '')

# ================================
# CLI INTERFACE
# ================================

async def main():
    """Main CLI interface"""
    import argparse
    
    parser = argparse.ArgumentParser(
        description='Enhanced Intelligent FSD Generator with Word Template Support',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Process single file with default template
  python main.py single input.html
  
  # Process single file with custom template
  python main.py single input.html --template /path/to/template.docx
  
  # Process single file with custom output directory
  python main.py single input.html --output-dir /path/to/output
  
  # Process multiple files
  python main.py batch file1.html file2.html file3.html --template /path/to/template.docx
  
  # Use custom configuration file
  python main.py single input.html --config config.json

Environment Variables:
  GEMINI_API_KEY      - Required: Gemini API key
  GEMINI_API_URL      - Optional: Custom API URL
  PROJECT_NAME        - Optional: Project name for FSD documents
  OUTPUT_DIR          - Optional: Default output directory
  TEMPLATE_DIR        - Optional: Template directory path
  MAX_TOKENS          - Optional: Maximum tokens for LLM responses (default: 4096)
  TEMPERATURE         - Optional: LLM temperature (default: 0.1)
        """
    )
    
    subparsers = parser.add_subparsers(dest='command', help='Available commands')
    
    # Single file command
    single_parser = subparsers.add_parser('single', help='Process single HTML file')
    single_parser.add_argument('html_file', help='Path to HTML file')
    single_parser.add_argument('--template', help='Word template file path')
    single_parser.add_argument('--output-dir', help='Custom output directory')
    single_parser.add_argument('--config', help='Configuration file path')
    
    # Batch processing command
    batch_parser = subparsers.add_parser('batch', help='Process multiple HTML files')
    batch_parser.add_argument('html_files', nargs='+', help='Paths to HTML files')
    batch_parser.add_argument('--template', help='Word template file path')
    batch_parser.add_argument('--output-dir', help='Custom output directory')
    batch_parser.add_argument('--config', help='Configuration file path')
    
    # Configuration command
    config_parser = subparsers.add_parser('config', help='Generate sample configuration file')
    config_parser.add_argument('--output', default='config.json', help='Output configuration file')
    
    args = parser.parse_args()
    
    if args.command == 'config':
        # Generate sample configuration
        sample_config = {
            "GEMINI_API_KEY": "AIzaSyD1S4r1Xz-7E2f8RyQFEehmzuHa7ZrINMM",
            "gemini_api_url": "https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-pro-latest:generateContent",
            "project_name": "System Integrator for Management Information System (MIS) Towards Single Source of Truth (SSoT)",
            "default_output_dir": "./output",
            "template_dir": "./Template/FSD",
            "max_tokens": 4096,
            "temperature": 0.1
        }
        
        with open(args.output, 'w', encoding='utf-8') as f:
            json.dump(sample_config, f, indent=2)
        
        print(f"Sample configuration saved to: {args.output}")
        print("Please update the configuration with your actual values.")
        return
    
    if not args.command:
        parser.print_help()
        return
    
    try:
        # Initialize generator
        generator = EnhancedIntelligentFSDGenerator(getattr(args, 'config', None))
        
        if args.command == 'single':
            # Process single file
            results = await generator.process_file(
                args.html_file, 
                getattr(args, 'template', None), 
                getattr(args, 'output_dir', None)
            )
            
            print("\n=== Processing Results ===")
            summary = results['analysis_summary']
            print(f"Program: {summary['program_name']}")
            print(f"Description: {summary['description']}")
            print(f"Selection Parameters: {summary['selection_parameters_count']}")
            print(f"Field Mappings: {summary['field_mappings_count']}")
            print(f"Error Scenarios: {summary['error_scenarios_count']}")
            print(f"Test Scenarios: {summary['test_scenarios_count']}")
            
            print("\n=== Generated Files ===")
            for output_type, file_path in results['output_files'].items():
                print(f"{output_type.upper()}: {file_path}")
        
        elif args.command == 'batch':
            # Process multiple files
            results = await generator.process_multiple_files(
                args.html_files, 
                getattr(args, 'template', None), 
                getattr(args, 'output_dir', None)
            )
            
            print(f"\n=== Batch Processing Results ===")
            print(f"Total files: {results['total_files']}")
            print(f"Successful: {results['successful']}")
            print(f"Failed: {results['failed']}")
            
            if results['failed'] > 0:
                print("\nFailed files:")
                for file_path, result in results['results'].items():
                    if 'error' in result:
                        print(f"  - {os.path.basename(file_path)}: {result['error']}")
    
    except Exception as e:
        logger.error(f"Error: {e}")
        return 1
    
    return 0

# ================================
# UTILITY FUNCTIONS
# ================================

def setup_requirements():
    """Setup requirements for the enhanced FSD generator"""
    requirements = [
        "aiohttp>=3.8.0",
        "beautifulsoup4>=4.11.0",
        "python-docx>=0.8.11",
        "docxtpl>=0.16.0",
        "markdown>=3.4.0",
        "markdown2>=2.4.0",
        "pathlib>=1.0.1",
        "fastapi>=0.100.0",
        "uvicorn>=0.22.0",
        "python-multipart>=0.0.6",
        "openpyxl>=3.1.0"
    ]
    
    print("Required packages:")
    for req in requirements:
        print(f"  - {req}")
    
    print("\nInstall with:")
    print("pip install " + " ".join(requirements))

def create_sample_config():
    """Create a sample configuration file"""
    config = {
        "GEMINI_API_KEY": "AIzaSyD1S4r1Xz-7E2f8RyQFEehmzuHa7ZrINMM",
        "gemini_api_url": "https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-pro-latest:generateContent",
        "project_name": "System Integrator for Management Information System (MIS) Towards Single Source of Truth (SSoT)",
        "default_output_dir": "./output",
        "template_dir": "./Template/FSD",
        "max_tokens": 4096,
        "temperature": 0.1,
        "default_template_path": r"C:\Users\wahyu.perwira\Documents\Project\poc\SAP-AUTOMATE-FD-TD\backend\templates\Template_PLN_SI SSoT_(DAPI ID)_(Module Name)_Functional Specification Design (FSD)_v100_ID.docx"
    }
    
    with open('fsd_config.json', 'w', encoding='utf-8') as f:
        json.dump(config, f, indent=2)
    
    print("Sample configuration saved to: fsd_config.json")
    print("Please update with your actual API key and paths.")

# Example usage functions
async def example_single_file():
    """Example: Process a single file"""
    config_file = "fsd_config.json"
    html_file = "sample_input.html"
    template_file = r"C:\Users\wahyu.perwira\Documents\Project\poc\SAP-AUTOMATE-FD-TD\backend\templates\Template_PLN_SI SSoT_(DAPI ID)_(Module Name)_Functional Specification Design (FSD)_v100_ID.docx"
    
    try:
        generator = EnhancedIntelligentFSDGenerator(config_file)
        results = await generator.process_file(html_file, template_file)
        
        print("Processing completed successfully!")
        print(f"Generated files:")
        for output_type, file_path in results['output_files'].items():
            print(f"  {output_type}: {file_path}")
            
    except Exception as e:
        print(f"Error: {e}")

async def example_batch_processing():
    """Example: Process multiple files"""
    config_file = "fsd_config.json"
    html_files = [
        "file1.html",
        "file2.html", 
        "file3.html"
    ]
    template_file = "/Users/wahyu.perwira/Documents/Project/poc/SAP-AUTOMATE-FD-TD/backend/output/templates/Template_PLN_SI SSoT_(DAPI ID)_(Module Name)_Functional Specification Design (FSD)_v100_ID.docx"
    
    try:
        generator = EnhancedIntelligentFSDGenerator(config_file)
        results = await generator.process_multiple_files(html_files, template_file)
        
        print(f"Batch processing completed!")
        print(f"Processed: {results['successful']}/{results['total_files']} files")
        
    except Exception as e:
        print(f"Error: {e}")

# ================================
# MAIN ENTRY POINT
# ================================

if __name__ == "__main__":
    import uvicorn
    
    # Ensure all directories exist
    for directory in [UPLOAD_DIR, OUTPUT_DIR, TEMPLATE_DIR, CONFIG_DIR]:
        os.makedirs(directory, exist_ok=True)
    
    logger.info("Starting Accenture SAP FSD Document Processor API")
    logger.info(f"FSD Generator available: {fsd_generator is not None}")
    
    # Check if running with CLI arguments
    import sys
    if len(sys.argv) > 1:
        # CLI mode
        asyncio.run(main())
    else:
        # FastAPI server mode
        uvicorn.run(
            app, 
            host="0.0.0.0", 
            port=8000, 
            reload=True,
            log_level="info"
        )